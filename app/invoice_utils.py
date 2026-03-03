import re
from dataclasses import dataclass
from typing import Dict

from docx import Document


def group4(s: str) -> str:
    digits = re.sub(r"\D+", "", s or "")
    return " ".join(digits[i:i + 4] for i in range(0, len(digits), 4)).strip()


def normalize_digits(s: str) -> str:
    return re.sub(r"\D+", "", s or "")


def pick_shorter_inn(*inn_values: str) -> str:
    inns = [normalize_digits(x) for x in inn_values if normalize_digits(x)]
    if not inns:
        return ""
    inns.sort(key=lambda x: (len(x), x))
    return inns[0]


def format_amount_ru(value: str) -> str:
    """Return amount formatted like 396 580,00 (ru style)."""
    s = (value or "").strip().replace("\u00A0", "").replace(" ", "").replace(",", ".")
    if not s:
        raise ValueError("empty amount")
    num = float(s)
    rub = int(num)
    kop = int(round((num - rub) * 100))
    if kop == 100:
        rub += 1
        kop = 0
    rub_str = f"{rub:,}".replace(",", " ")
    return f"{rub_str},{kop:02d}"


@dataclass
class Requisites:
    recipient_fio: str = ""
    bank_name: str = ""
    bank_inn: str = ""
    bank_kpp: str = ""
    bank_bik: str = ""
    bank_ks: str = ""
    bank_rs: str = ""


def parse_requisites(text: str) -> Requisites:
    """Parse requisites from pasted block (taken from provided bot logic)."""
    t = (text or "").replace("\u00A0", " ")

    def find(pattern: str) -> str:
        m = re.search(pattern, t, re.IGNORECASE | re.MULTILINE)
        return m.group(1).strip() if m else ""

    recipient_fio = find(r"Получатель\s*[:\-]?\s*([А-ЯЁA-Z][А-ЯЁA-Z\s]+)")
    bank_name = find(r"(?:Банк получателя|Банк)\s*[:\-]?\s*(.+)")
    bank_bik = find(r"БИК(?:\s+банка\s+получателя)?\s*[:\-]?\s*(\d{9})")
    bank_kpp = find(r"КПП\s*[:\-]?\s*(\d{9})")

    bank_rs = find(r"(?:Сч[её]т|Номер\s*сч[её]та)\s*(?:получателя)?\s*[:\-]?\s*(\d{20})")
    bank_ks = find(r"(?:Корр\.?\s*сч[её]т|Кор\.?\s*сч[её]т|K\/C)\s*[:\-]?\s*(\d{20})")

    inns = re.findall(r"ИНН\s*[:\-]?\s*(\d{10,12})", t, flags=re.IGNORECASE)
    bank_inn = pick_shorter_inn(*inns)

    return Requisites(
        recipient_fio=recipient_fio,
        bank_name=bank_name,
        bank_inn=bank_inn,
        bank_kpp=bank_kpp,
        bank_bik=bank_bik,
        bank_ks=group4(bank_ks),
        bank_rs=group4(bank_rs),
    )


def replace_placeholders(doc: Document, mapping: Dict[str, str]) -> None:
    """Replace placeholders in docx robustly across runs/tables."""

    def replace_in_paragraph(p):
        if not p.runs:
            return
        full_text = "".join(r.text for r in p.runs)
        new_text = full_text
        for k, v in mapping.items():
            new_text = new_text.replace(k, v)
        if new_text == full_text:
            return
        for run in p.runs:
            run.text = ""
        p.runs[0].text = new_text

    for p in doc.paragraphs:
        replace_in_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

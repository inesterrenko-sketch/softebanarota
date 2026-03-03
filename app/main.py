import json
import os
import re
import shutil
import sqlite3
import subprocess
import tempfile
import uuid
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

from docx import Document
from docx.shared import Inches
from num2words import num2words

try:
    import requests
    REQUESTS_AVAILABLE = True
except ImportError:
    REQUESTS_AVAILABLE = False

BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
UPLOADS_DIR = DATA_DIR / "uploads"
TEMPLATES_DB = DATA_DIR / "templates.json"
GENERATED_DIR = DATA_DIR / "generated"
CLIENTS_DB = DATA_DIR / "clients.db"
INVOICE_TEMPLATE_PATH = DATA_DIR / "invoice_template.docx"
TPO_TEMPLATE_PATH = DATA_DIR / "tpo_template.docx"
KVIT_TEMPLATE_PATH = DATA_DIR / "kvit_template.docx"
UTIL_TEMPLATE_PATH = DATA_DIR / "util_template.docx"
DKP_TEMPLATE_PATH = DATA_DIR / "dkp_template.docx"
PROXY_TEMPLATE_PATH = DATA_DIR / "proxy_template.docx"
SETTINGS_PATH = DATA_DIR / "settings.json"
COMPANIES_PATH = DATA_DIR / "companies.json"

# Таблица транслитерации русский -> латиница
TRANSLIT_MAP = {
    'а': 'a', 'б': 'b', 'в': 'v', 'г': 'g', 'д': 'd', 'е': 'e', 'ё': 'e',
    'ж': 'zh', 'з': 'z', 'и': 'i', 'й': 'y', 'к': 'k', 'л': 'l', 'м': 'm',
    'н': 'n', 'о': 'o', 'п': 'p', 'р': 'r', 'с': 's', 'т': 't', 'у': 'u',
    'ф': 'f', 'х': 'kh', 'ц': 'ts', 'ч': 'ch', 'ш': 'sh', 'щ': 'shch',
    'ъ': '', 'ы': 'y', 'ь': '', 'э': 'e', 'ю': 'yu', 'я': 'ya',
    'А': 'A', 'Б': 'B', 'В': 'V', 'Г': 'G', 'Д': 'D', 'Е': 'E', 'Ё': 'E',
    'Ж': 'Zh', 'З': 'Z', 'И': 'I', 'Й': 'Y', 'К': 'K', 'Л': 'L', 'М': 'M',
    'Н': 'N', 'О': 'O', 'П': 'P', 'Р': 'R', 'С': 'S', 'Т': 'T', 'У': 'U',
    'Ф': 'F', 'Х': 'Kh', 'Ц': 'Ts', 'Ч': 'Ch', 'Ш': 'Sh', 'Щ': 'Shch',
    'Ъ': '', 'Ы': 'Y', 'Ь': '', 'Э': 'E', 'Ю': 'Yu', 'Я': 'Ya',
}

# Перевод цветов на немецкий
COLOR_TRANSLATION = {
    'белый': 'weiß',
    'черный': 'schwarz',
    'серебристый': 'silber',
    'серый': 'grau',
    'синий': 'blau',
    'голубой': 'hellblau',
    'зеленый': 'grün',
    'красный': 'rot',
    'оранжевый': 'orange',
    'желтый': 'gelb',
    'коричневый': 'braun',
    'бежевый': 'beige',
    'бордовый': 'burgunder',
    'фиолетовый': 'violett',
    'розовый': 'rosa',
    'золотистый': 'gold',
}

SEED_CLIENTS_DB = Path("/mnt/data/clients.db")
SEED_INVOICE_TEMPLATE = Path("/mnt/data/template.docx")
SEED_TPO_TEMPLATE = Path("/mnt/data/template.docx")

SAMPLE_TEMPLATES = [
    {
        "company": "Автоберг",
        "name": "Договор подбора авто",
        "source_path": "/mnt/data/Автоберг.docx",
    },
    {
        "company": "Автолюкс",
        "name": "Договор подбора авто",
        "source_path": "/mnt/data/автолюкс.docx",
    },
    {
        "company": "АТЦ Белгород",
        "name": "Договор подбора авто",
        "source_path": "/mnt/data/атц белгород.docx",
    },
]

FIELD_PRESETS: Dict[str, Dict[str, str]] = {
    "CONTRACT_NO": {"label": "Номер договора", "type": "text"},
    "CONTRACT_DATE": {"label": "Дата договора", "type": "date"},
    "CLIENT_FIO": {"label": "ФИО клиента", "type": "text"},
    "PASSPORT": {"label": "Паспорт (серия и номер)", "type": "text"},
    "VYDAN": {"label": "Дата выдачи паспорта", "type": "date"},
    "ORGAN": {"label": "Кем выдан паспорт", "type": "textarea"},
    "ADDRESS": {"label": "Адрес", "type": "textarea"},
    "PHONE": {"label": "Телефон", "type": "tel"},
    "CAR": {"label": "Марка и модель", "type": "text"},
    "year car": {"label": "Год выпуска", "type": "number"},
    "vin": {"label": "VIN", "type": "text"},
    "Probeg": {"label": "Пробег (км)", "type": "number"},
    "OBEM": {"label": "Объем двигателя (см³)", "type": "number"},
    "Engine": {"label": "Тип двигателя / топливо", "type": "text"},
    "Color": {"label": "Цвет", "type": "text"},
    "CUST_RUB": {"label": "Растаможка (₽)", "type": "number"},
    "CAR_RUB": {"label": "Стоимость авто (₽)", "type": "number"},
    "CAR_ALL": {"label": "Итого под ключ (₽)", "type": "number"},
    "Delivery": {"label": "Адрес доставки", "type": "textarea"},
}

MANDATORY_CONTRACT_FIELDS = ["CLIENT_FIO", "PASSPORT", "ORGAN", "VYDAN", "ADDRESS", "PHONE", "CONTRACT_NO", "CONTRACT_DATE"]

PLACEHOLDER_RE = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")


def ensure_dirs() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)


def ensure_seed_files() -> None:
    try:
        if (not CLIENTS_DB.exists()) and SEED_CLIENTS_DB.exists():
            shutil.copy2(SEED_CLIENTS_DB, CLIENTS_DB)
        if (not INVOICE_TEMPLATE_PATH.exists()) and SEED_INVOICE_TEMPLATE.exists():
            shutil.copy2(SEED_INVOICE_TEMPLATE, INVOICE_TEMPLATE_PATH)
        if (not TPO_TEMPLATE_PATH.exists()) and SEED_TPO_TEMPLATE.exists():
            shutil.copy2(SEED_TPO_TEMPLATE, TPO_TEMPLATE_PATH)
    except Exception:
        pass


def db_connect() -> sqlite3.Connection:
    con = sqlite3.connect(str(CLIENTS_DB))
    con.row_factory = sqlite3.Row
    return con


def db_init() -> None:
    con = db_connect()
    cur = con.cursor()
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS clients (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            fio TEXT NOT NULL DEFAULT '',
            contract_no TEXT NOT NULL DEFAULT '',
            contract_date TEXT NOT NULL DEFAULT '',
            car_model TEXT NOT NULL DEFAULT ''
        )
        """
    )
    con.commit()

    cur.execute("PRAGMA table_info(clients)")
    existing = {r[1] for r in cur.fetchall()}
    additions: List[Tuple[str, str]] = [
        ("passport", "TEXT NOT NULL DEFAULT ''"),
        ("organ", "TEXT NOT NULL DEFAULT ''"),
        ("vydan", "TEXT NOT NULL DEFAULT ''"),
        ("address", "TEXT NOT NULL DEFAULT ''"),
        ("phone", "TEXT NOT NULL DEFAULT ''"),
        ("customs_amount", "TEXT NOT NULL DEFAULT ''"),
        ("dkp_amount", "TEXT NOT NULL DEFAULT ''"),
        ("last_template_id", "TEXT NOT NULL DEFAULT ''"),
        ("last_contract_json", "TEXT NOT NULL DEFAULT ''"),
        ("vin", "TEXT NOT NULL DEFAULT ''"),
        ("obem", "TEXT NOT NULL DEFAULT ''"),
        ("vypusk", "TEXT NOT NULL DEFAULT ''"),
        ("registr", "TEXT NOT NULL DEFAULT ''"),
        ("last_tpo_json", "TEXT NOT NULL DEFAULT ''"),
        ("tpo_summa_byn", "TEXT NOT NULL DEFAULT ''"),
        ("tpo_price_eur", "TEXT NOT NULL DEFAULT ''"),
        ("company_inn", "TEXT NOT NULL DEFAULT ''"),
        ("company_address", "TEXT NOT NULL DEFAULT ''"),
        ("company_name", "TEXT NOT NULL DEFAULT ''"),
        ("company_id", "TEXT NOT NULL DEFAULT ''"),
        ("updated_at", "TEXT NOT NULL DEFAULT ''"),
        ("color", "TEXT NOT NULL DEFAULT ''"),
        ("probeg", "TEXT NOT NULL DEFAULT ''"),
        ("engine", "TEXT NOT NULL DEFAULT ''"),
        ("delivery", "TEXT NOT NULL DEFAULT ''"),
    ]
    for col, decl in additions:
        if col not in existing:
            cur.execute(f"ALTER TABLE clients ADD COLUMN {col} {decl}")
    con.commit()
    con.close()


def db_list_clients(limit: int = 200) -> List[Dict[str, str]]:
    con = db_connect()
    cur = con.cursor()
    cur.execute("SELECT * FROM clients ORDER BY id DESC LIMIT ?", (limit,))
    rows = cur.fetchall()
    con.close()
    return [dict(r) for r in rows]


def db_get_client(cid: int) -> Optional[Dict[str, str]]:
    con = db_connect()
    cur = con.cursor()
    cur.execute("SELECT * FROM clients WHERE id=?", (cid,))
    row = cur.fetchone()
    con.close()
    return dict(row) if row else None


def db_upsert_client(cid: Optional[int], fields: Dict[str, str]) -> int:
    now = datetime.utcnow().isoformat() + "Z"
    fields = {k: (v or "") for k, v in fields.items()}
    con = db_connect()
    cur = con.cursor()
    if cid:
        cols = sorted(fields.keys())
        sets = ", ".join([f"{c}=?" for c in cols] + ["updated_at=?"])
        values = [fields[c] for c in cols] + [now, cid]
        cur.execute(f"UPDATE clients SET {sets} WHERE id=?", values)
        con.commit()
        con.close()
        return int(cid)

    cols = sorted(fields.keys())
    cols_sql = ", ".join(cols + ["updated_at"])
    qs = ", ".join(["?"] * (len(cols) + 1))
    values = [fields[c] for c in cols] + [now]
    cur.execute(f"INSERT INTO clients ({cols_sql}) VALUES ({qs})", values)
    con.commit()
    new_id = int(cur.lastrowid)
    con.close()
    return new_id


def db_delete_client(cid: int) -> bool:
    """Удалить клиента и все связанные данные"""
    con = db_connect()
    cur = con.cursor()
    
    # Удаляем связанные записи в правильном порядке (учитывая внешние ключи)
    cur.execute("DELETE FROM proxy WHERE client_id=?", (cid,))
    cur.execute("DELETE FROM invoice_history WHERE client_id=?", (cid,))
    
    # Удаляем клиента
    cur.execute("DELETE FROM clients WHERE id=?", (cid,))
    
    con.commit()
    deleted = cur.rowcount > 0
    con.close()
    return deleted


def db_clear_client_history(cid: int) -> bool:
    """Очистить историю документов клиента"""
    con = db_connect()
    cur = con.cursor()
    
    cur.execute("DELETE FROM invoice_history WHERE client_id=?", (cid,))
    cur.execute("DELETE FROM proxy WHERE client_id=?", (cid,))
    cur.execute("DELETE FROM contracts WHERE client_id=?", (cid,))
    
    con.commit()
    con.close()
    return True


def translit_ru_to_lat(text: str) -> str:
    if not text:
        return text
    result = []
    for char in text:
        result.append(TRANSLIT_MAP.get(char, char))
    return ''.join(result)


def translate_ru_to_de(text: str) -> str:
    if not text or not REQUESTS_AVAILABLE:
        return text
    
    # Сначала пробуем перевести через API
    try:
        api_url = f"https://api.mymemory.translated.net/get?q={text}&langpair=ru|de"
        response = requests.get(api_url, timeout=10)
        
        if response.status_code == 200:
            data = response.json()
            translated = data.get('responseData', {}).get('translatedText', '')
            if translated and translated != text:
                return translated
    except Exception:
        pass
    
    # Если API недоступен, используем словарь для цветов
    text_lower = text.lower().strip()
    if text_lower in COLOR_TRANSLATION:
        return COLOR_TRANSLATION[text_lower]
    
    return text


def translate_color_ru_to_de(color_text: str) -> str:
    """Переводит цвет с русского на немецкий"""
    if not color_text:
        return ""
    
    color_lower = color_text.lower().strip()
    
    # Проверяем, может это уже немецкий цвет
    german_colors = ['weiß', 'schwarz', 'silber', 'grau', 'blau', 'hellblau', 
                     'grün', 'rot', 'orange', 'gelb', 'braun', 'beige', 
                     'burgunder', 'violett', 'rosa', 'gold']
    
    if color_lower in german_colors:
        return color_text  # Уже на немецком
    
    # Переводим с русского на немецкий
    if color_lower in COLOR_TRANSLATION:
        return COLOR_TRANSLATION[color_lower]
    
    # Если не нашли в словаре, возвращаем как есть
    return color_text


def convert_date_to_mmyyyy(date_str: str) -> str:
    """Преобразует дату из дд.мм.гггг в мм.гггг"""
    if not date_str:
        return ""
    
    # Если уже в формате мм.гггг, возвращаем как есть
    if re.match(r'^\d{2}\.\d{4}$', date_str):
        return date_str
    
    # Если в формате дд.мм.гггг, преобразуем
    match = re.match(r'^\d{2}\.(\d{2})\.(\d{4})$', date_str)
    if match:
        month, year = match.groups()
        return f"{month}.{year}"
    
    return date_str


def decline_fio_genitive(fio: str) -> str:
    """Склонение ФИО в родительный падеж (кого?) - Иванова Ивана Ивановича"""
    if not fio:
        return ""
    
    parts = fio.strip().split()
    if len(parts) < 2:
        return fio
    
    surname, name = parts[0], parts[1]
    patronymic = parts[2] if len(parts) > 2 else ""
    
    # Фамилия
    if surname.endswith('ов') or surname.endswith('ев') or surname.endswith('ин') or surname.endswith('ын'):
        surname_gen = surname + 'а'
    elif surname.endswith('ова') or surname.endswith('ева') or surname.endswith('ина') or surname.endswith('ына'):
        surname_gen = surname[:-1] + 'ой'
    elif surname.endswith('ая'):
        surname_gen = surname[:-2] + 'ой'
    elif surname.endswith('ий'):
        surname_gen = surname[:-2] + 'ого'
    elif surname.endswith('ый'):
        surname_gen = surname[:-2] + 'ого'
    else:
        surname_gen = surname
    
    # Имя
    if name.endswith('а') or name.endswith('я'):
        # Женское имя
        if name.endswith('ия'):
            name_gen = name[:-1] + 'и'
        else:
            name_gen = name[:-1] + 'ы'
    else:
        # Мужское имя
        if name.endswith('й'):
            name_gen = name[:-1] + 'я'
        elif name.endswith('ь'):
            name_gen = name[:-1] + 'я'
        else:
            name_gen = name + 'а'
    
    # Отчество
    if patronymic:
        if patronymic.endswith('вна') or patronymic.endswith('чна'):
            patronymic_gen = patronymic[:-1] + 'ы'
        elif patronymic.endswith('ович') or patronymic.endswith('евич'):
            patronymic_gen = patronymic + 'а'
        else:
            patronymic_gen = patronymic
        
        return f"{surname_gen} {name_gen} {patronymic_gen}"
    
    return f"{surname_gen} {name_gen}"


def decline_fio_instrumental(fio: str) -> str:
    """Склонение ФИО в творительный падеж (кем?) - Ивановым Иваном Ивановичем"""
    if not fio:
        return ""
    
    parts = fio.strip().split()
    if len(parts) < 2:
        return fio
    
    surname, name = parts[0], parts[1]
    patronymic = parts[2] if len(parts) > 2 else ""
    
    # Фамилия
    if surname.endswith('ов') or surname.endswith('ев') or surname.endswith('ин') or surname.endswith('ын'):
        surname_instr = surname + 'ым'
    elif surname.endswith('ова') or surname.endswith('ева') or surname.endswith('ина') or surname.endswith('ына'):
        surname_instr = surname[:-1] + 'ой'
    elif surname.endswith('ая'):
        surname_instr = surname[:-2] + 'ой'
    elif surname.endswith('ий'):
        surname_instr = surname[:-2] + 'им'
    elif surname.endswith('ый'):
        surname_instr = surname[:-2] + 'ым'
    else:
        surname_instr = surname + 'ом'
    
    # Имя
    if name.endswith('а') or name.endswith('я'):
        # Женское имя
        if name.endswith('ия'):
            name_instr = name[:-1] + 'ей'
        else:
            name_instr = name[:-1] + 'ой'
    else:
        # Мужское имя
        if name.endswith('й'):
            name_instr = name[:-1] + 'ем'
        elif name.endswith('ь'):
            name_instr = name[:-1] + 'ем'
        else:
            name_instr = name + 'ом'
    
    # Отчество
    if patronymic:
        if patronymic.endswith('вна') or patronymic.endswith('чна'):
            patronymic_instr = patronymic[:-1] + 'ой'
        elif patronymic.endswith('ович'):
            patronymic_instr = patronymic + 'ем'
        elif patronymic.endswith('евич'):
            patronymic_instr = patronymic + 'ем'
        else:
            patronymic_instr = patronymic
        
        return f"{surname_instr} {name_instr} {patronymic_instr}"
    
    return f"{surname_instr} {name_instr}"


def decline_fio_accusative(fio: str) -> str:
    """Склонение ФИО в винительный падеж (кого?) - Иванову Ольгу Сергеевну / того же Ивана"""
    if not fio:
        return ""
    
    parts = fio.strip().split()
    if len(parts) < 2:
        return fio
    
    surname, name = parts[0], parts[1]
    patronymic = parts[2] if len(parts) > 2 else ""
    
    # Определяем пол по отчеству или имени
    is_female = False
    if patronymic and (patronymic.endswith('вна') or patronymic.endswith('чна')):
        is_female = True
    elif name.endswith('а') or name.endswith('я'):
        is_female = True
    
    # Фамилия
    if is_female:
        if surname.endswith('ова') or surname.endswith('ева') or surname.endswith('ина') or surname.endswith('ына'):
            surname_acc = surname[:-1] + 'у'
        elif surname.endswith('ая'):
            surname_acc = surname[:-2] + 'ую'
        else:
            surname_acc = surname
    else:
        # Мужская фамилия не изменяется в винительном падеже
        surname_acc = surname
    
    # Имя
    if is_female:
        if name.endswith('ия'):
            name_acc = name[:-1] + 'ю'
        elif name.endswith('а') or name.endswith('я'):
            name_acc = name[:-1] + 'у'
        else:
            name_acc = name
    else:
        # Мужское имя = родительный = генитив
        if name.endswith('й'):
            name_acc = name[:-1] + 'я'
        elif name.endswith('ь'):
            name_acc = name[:-1] + 'я'
        else:
            name_acc = name + 'а'
    
    # Отчество
    if patronymic:
        if patronymic.endswith('вна') or patronymic.endswith('чна'):
            patronymic_acc = patronymic[:-1] + 'у'
        elif patronymic.endswith('ович') or patronymic.endswith('евич'):
            patronymic_acc = patronymic + 'а'
        else:
            patronymic_acc = patronymic
        
        return f"{surname_acc} {name_acc} {patronymic_acc}"
    
    return f"{surname_acc} {name_acc}"


def format_passport_dkp(text: str) -> str:
    digits = re.sub(r'\D', '', text or '')
    if len(digits) < 6:
        return text
    return f"{digits[:2]} {digits[2:4]} {digits[4:10]}"


def format_phone_dkp(text: str) -> str:
    digits = re.sub(r'\D', '', text or '')
    
    if not digits:
        return text
    
    if len(digits) == 11:
        if digits[0] == '8':
            digits = '7' + digits[1:]
        elif digits[0] != '7':
            return text
    elif len(digits) == 10:
        digits = '7' + digits
    else:
        return text
    
    return f"+{digits[0]} {digits[1:4]} {digits[4:7]} {digits[7:11]}"


def format_vin_dkp(text: str) -> str:
    return (text or '').upper().replace('-', '').replace(' ', '').strip()


def format_regdate_dkp(text: str) -> str:
    match = re.match(r'(\d{1,2})[./\-\s](\d{4})', text or '')
    if match:
        month, year = match.groups()
        return f"{int(month):02d}.{year}"
    
    match = re.match(r'\d{1,2}[./\-\s](\d{1,2})[./\-\s](\d{4})', text or '')
    if match:
        month, year = match.groups()
        return f"{int(month):02d}.{year}"
    
    return text or ''


def ps_to_kw(ps: str) -> str:
    try:
        kw = float(ps) * 0.735499
        return str(round(kw, 1))
    except (ValueError, TypeError):
        return ''


def group4(s: str) -> str:
    digits = re.sub(r"\D+", "", s or "")
    return " ".join(digits[i:i + 4] for i in range(0, len(digits), 4)).strip()


def normalize_digits(s: str) -> str:
    return re.sub(r"\D+", "", s or "")


def pick_shorter_inn(*inn_values: str) -> str:
    inns = [normalize_digits(x) for x in inn_values if normalize_digits(x)]
    if not inns:
        return ""
    inns.sort(key=lambda x: (len(x), x))
    return inns[0]


def format_amount_ru(value: str) -> str:
    s = (value or "").strip().replace("\u00A0", "").replace(" ", "").replace(",", ".")
    if not s:
        raise ValueError("empty amount")
    num = float(s)
    rub = int(num)
    kop = int(round((num - rub) * 100))
    if kop == 100:
        rub += 1
        kop = 0
    rub_str = f"{rub:,}".replace(",", " ")
    return f"{rub_str},{kop:02d}"


@dataclass
class Requisites:
    recipient_fio: str = ""
    bank_name: str = ""
    bank_inn: str = ""
    bank_kpp: str = ""
    bank_bik: str = ""
    bank_ks: str = ""
    bank_rs: str = ""


def parse_requisites(text: str) -> Requisites:
    t = (text or "").replace("\u00A0", " ")

    def find(pattern: str) -> str:
        m = re.search(pattern, t, re.IGNORECASE | re.MULTILINE)
        return m.group(1).strip() if m else ""

    recipient_fio_match = re.search(r"Получатель\s*[:\-]?\s*([А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+)", t, re.IGNORECASE | re.MULTILINE)
    recipient_fio = recipient_fio_match.group(1).strip() if recipient_fio_match else ""
    bank_name = find(r"(?:Банк получателя-Банк)\s*[:\-]?\s*(.+)")
    bank_bik = find(r"БИК(?:\s+банка\s+получателя)?\s*[:\-]?\s*(\d{9})")
    bank_kpp = find(r"КПП\s*[:\-]?\s*(\d{9})")

    bank_rs = find(r"Номер\s+сч[её]та\s+получателя\s*[:\-]?\s*(\d{20})")
    if not bank_rs:
        bank_rs = find(r"Номер\s+сч[её]та\s*[:\-]?\s*(\d{20})")
    bank_ks = find(r"(?:Корр?\.?\s*сч[её]т[а-яА-Я]*|K\/C)\s*[:\-]?\s*(\d{20})")

    inns = re.findall(r"ИНН\s*[:\-]?\s*(\d{10,12})", t, flags=re.IGNORECASE)
    inns_12 = [x for x in inns if len(x) == 12]
    bank_inn = inns_12[0] if inns_12 else pick_shorter_inn(*inns)

    return Requisites(
        recipient_fio=recipient_fio,
        bank_name=bank_name,
        bank_inn=bank_inn,
        bank_kpp=bank_kpp,
        bank_bik=bank_bik,
        bank_ks=group4(bank_ks),
        bank_rs=group4(bank_rs),
    )


def _find_soffice() -> Optional[str]:
    if shutil.which("soffice"):
        return "soffice"
    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    return None


def docx_to_pdf(input_docx: str, output_pdf: str) -> None:
    outdir = os.path.dirname(output_pdf)
    os.makedirs(outdir, exist_ok=True)

    soffice = _find_soffice()
    if not soffice:
        raise RuntimeError("LibreOffice (soffice) не найден. Установи LibreOffice или добавь soffice.exe в PATH.")

    subprocess.run(
        [soffice, "--headless", "--nologo", "--convert-to", "pdf", "--outdir", outdir, input_docx],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )

    generated = os.path.join(outdir, os.path.splitext(os.path.basename(input_docx))[0] + ".pdf")
    if not os.path.exists(generated):
        raise RuntimeError("PDF не был создан LibreOffice.")
    shutil.move(generated, output_pdf)


def iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
        for p in section.footer.paragraphs:
            yield p
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p


def extract_placeholders(doc: Document) -> List[str]:
    found: Set[str] = set()
    for p in iter_all_paragraphs(doc):
        text = p.text or ""
        for m in PLACEHOLDER_RE.findall(text):
            key = m.strip()
            if key:
                found.add(key)
    return sorted(found)


def replace_in_paragraph(p, mapping: Dict[str, str]) -> None:
    if not p.runs:
        return
    full = "".join(r.text for r in p.runs)
    new = full

    for key, value in mapping.items():
        new = new.replace(f"{{{{{key}}}}}", value)
        new = re.sub(r"\{\{\s*" + re.escape(key) + r"\s*\}\}", value, new)

    if new != full:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""


def apply_mapping(doc: Document, mapping: Dict[str, str]) -> None:
    for p in iter_all_paragraphs(doc):
        replace_in_paragraph(p, mapping)


def insert_qr_code(doc: Document, image_path: str, placeholder_key: str = "QR CODE") -> bool:
    token = f"{{{{{placeholder_key}}}}}"
    inserted = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if token in (cell.text or ""):
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.text = ""
                    
                    while len(cell.paragraphs) > 1:
                        try:
                            cell._element.remove(cell.paragraphs[-1]._element)
                        except:
                            break
                    
                    if not cell.paragraphs:
                        p = cell.add_paragraph()
                    else:
                        p = cell.paragraphs[0]
                    
                    try:
                        p.alignment = 1
                    except Exception:
                        pass
                    
                    run = p.add_run()
                    run.add_picture(image_path, width=Inches(1.3), height=Inches(1.3))
                    inserted = True
    return inserted


def load_db() -> Dict[str, Any]:
    if not TEMPLATES_DB.exists():
        return {"templates": []}
    return json.loads(TEMPLATES_DB.read_text(encoding="utf-8"))


def save_db(db: Dict[str, Any]) -> None:
    TEMPLATES_DB.write_text(json.dumps(db, ensure_ascii=False, indent=2), encoding="utf-8")


DEFAULT_SETTINGS: Dict[str, Any] = {
    "currency": {
        "rub_to_byn": 0.04,
        "eur_byn": 4.0,
        "usd_byn": 3.4,
    }
}


def load_settings() -> Dict[str, Any]:
    if not SETTINGS_PATH.exists():
        return json.loads(json.dumps(DEFAULT_SETTINGS))
    try:
        data = json.loads(SETTINGS_PATH.read_text(encoding="utf-8"))
        out = json.loads(json.dumps(DEFAULT_SETTINGS))
        out.update(data or {})
        out.setdefault("currency", {}).update((data or {}).get("currency", {}))
        return out
    except Exception:
        return json.loads(json.dumps(DEFAULT_SETTINGS))


def save_settings(data: Dict[str, Any]) -> None:
    SETTINGS_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def get_currency_rates() -> Dict[str, float]:
    s = load_settings().get("currency", {})
    def f(k: str, default: float) -> float:
        try:
            return float(str(s.get(k, default)).replace(",", ".").strip())
        except Exception:
            return float(default)
    return {
        "rub_to_byn": f("rub_to_byn", 0.04),
        "eur_byn": f("eur_byn", 4.0),
        "usd_byn": f("usd_byn", 3.4),
    }


def upsert_sample_templates() -> None:
    db = load_db()
    existing = {(t.get("company"), t.get("name")) for t in db.get("templates", [])}

    changed = False
    for t in SAMPLE_TEMPLATES:
        src = Path(t["source_path"])
        if not src.exists():
            continue
        key = (t["company"], t["name"])
        if key in existing:
            continue

        template_id = uuid.uuid4().hex
        dest_name = f"{template_id}.docx"
        dest = UPLOADS_DIR / dest_name
        shutil.copy2(src, dest)

        doc = Document(str(dest))
        placeholders = extract_placeholders(doc)

        db.setdefault("templates", []).append(
            {
                "id": template_id,
                "company": t["company"],
                "name": t["name"],
                "filename": dest_name,
                "placeholders": placeholders,
                "created_at": datetime.utcnow().isoformat() + "Z",
            }
        )
        changed = True

    if changed:
        save_db(db)


def get_template(template_id: str) -> Dict[str, Any]:
    db = load_db()
    for t in db.get("templates", []):
        if t.get("id") == template_id:
            return t
    raise HTTPException(status_code=404, detail="Template not found")


def list_templates() -> List[Dict[str, Any]]:
    db = load_db()
    return sorted(db.get("templates", []), key=lambda x: (x.get("company", ""), x.get("name", "")))


def safe_filename(s: str) -> str:
    s = s.strip()
    s = re.sub(r"\s+", "_", s)
    s = re.sub(r"[^0-9A-Za-zА-Яа-я_\-\.]+", "", s)
    return s[:80] if s else uuid.uuid4().hex


def normalize_date(value: str) -> str:
    try:
        dt = datetime.strptime(value, "%Y-%m-%d")
        return dt.strftime("%d.%m.%Y")
    except Exception:
        return value


def _pick_first(mapping: Dict[str, str], keys: List[str]) -> str:
    for k in keys:
        v = (mapping.get(k) or "").strip()
        if v:
            return v
    return ""


def _find_amount_like(mapping: Dict[str, str], substrings: List[str]) -> str:
    for k, v in mapping.items():
        ku = k.upper()
        if not any(s in ku for s in substrings):
            continue
        vv = (v or "").strip()
        if re.search(r"\d", vv):
            return vv
    return ""


ensure_dirs()
ensure_seed_files()
db_init()
upsert_sample_templates()

app = FastAPI(title="ContractFill")

@app.on_event("startup")
def startup_event():
    try:
        conn = db_connect()
        cur = conn.cursor()
        
        cur.execute("SELECT id FROM clients WHERE contract_no = ?", ("5266",))
        if not cur.fetchone():
            client_data = {
                'fio': 'Баженов Евгений Александрович',
                'passport': '76 25 415097',
                'organ': 'УМВД РОССИИ ПО ЗАБАЙКАЛЬСКОМУ КРАЮ',
                'vydan': '14.11.2025',
                'address': 'Кировская обл., М.Р-Н. Даровской, Г.П. Даровское, ПГТ. Даровское, ул. Зеленая, д.8 кв. 1',
                'phone': '+7 (912) 332-08-63',
                'contract_no': '5266',
                'contract_date': '16.01.2026',
                'car_model': 'Toyota RAV4',
                'vin': 'JTMW43FV80D135612',
                'obem': '1987',
                'vypusk': '2023',
                'customs_amount': '514445',
                'dkp_amount': '1145557',
                'company_inn': '2632083090',
                'company_address': '357204, Ставропольский край, Минераловоский р-н, тер. Автодорога, Р-217 Кавказ, км. 345-ый',
            }
            db_upsert_client(None, client_data)
            print("✅ Добавлен пример клиента: Баженов Е.А. (Договор №5266)")
        
        conn.close()
    except Exception as e:
        print(f"Ошибка при добавлении примера клиента: {e}")

app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")

jinja = Jinja2Templates(directory=str(BASE_DIR / "templates"))


@app.get("/")
def root():
    return RedirectResponse(url="/dashboard", status_code=303)


@app.get("/dashboard", response_class=HTMLResponse)
def dashboard(request: Request):
    templates = list_templates()
    clients = db_list_clients(200)
    companies = sorted({(t.get("company") or "").strip() for t in templates if (t.get("company") or "").strip()})
    currency = get_currency_rates()
    return jinja.TemplateResponse(
        "dashboard.html",
        {
            "request": request,
            "templates": templates,
            "clients": clients,
            "companies": companies,
            "currency": currency,
        },
    )


@app.get("/documents", response_class=HTMLResponse)
def documents_home(request: Request):
    templates = list_templates()
    clients = db_list_clients(200)
    companies = sorted({(t.get("company") or "").strip() for t in templates if (t.get("company") or "").strip()})
    return jinja.TemplateResponse(
        "home.html",
        {
            "request": request,
            "templates": templates,
            "clients": clients,
            "companies": companies,
        },
    )


@app.get("/api/clients")
def api_clients(limit: int = 200):
    return {"clients": db_list_clients(limit)}


@app.get("/api/clients/{cid}")
def api_client(cid: int):
    c = db_get_client(cid)
    if not c:
        raise HTTPException(status_code=404, detail="Client not found")
    return c


INVOICE_PURPOSES = [
    "Оплата таможенной пошлины за автомобиль {модель}",
    "Оплата по договору ДКП за автомобиль {модель}",
    "Оплата гос. пошлины за автомобиль {модель}",
    "Оплата утиль. сбора за автомобиль {модель}",
]


def _to_float(value: str) -> float:
    s = (value or "").replace("\u00A0", " ").strip()
    s = s.replace(" ", "").replace(",", ".")
    if not s:
        return 0.0
    try:
        return float(s)
    except Exception:
        s2 = re.sub(r"[^0-9\.]", "", s)
        try:
            return float(s2) if s2 else 0.0
        except Exception:
            return 0.0


def calc_tpo_values(price_rub: str, cust_rub: str, rates: Dict[str, float]) -> Dict[str, str]:
    rub_to_byn = float(rates.get("rub_to_byn", 0.04) or 0.04)
    eur_byn = float(rates.get("eur_byn", 4.0) or 4.0)
    usd_byn = float(rates.get("usd_byn", 3.4) or 3.4)

    pr = _to_float(price_rub)
    cu = _to_float(cust_rub)

    price_byn = pr * rub_to_byn
    cust_byn = cu * rub_to_byn

    price_eur = (price_byn / eur_byn) if eur_byn else 0.0
    price_usd = (price_byn / usd_byn) if usd_byn else 0.0

    percent = (cust_byn / price_byn * 100.0) if price_byn else 0.0

    return {
        "KURS": f"{eur_byn:.4f}",
        "PRICE": f"{price_byn:.2f}",
        "CUST": f"{cust_byn:.2f}",
        "PRICE_EUR": f"{price_eur:.2f}",
        "P_usd": f"{price_usd:.2f}",
        "PERCENT": f"{percent:.2f}%",
    }


@app.get("/invoice", response_class=HTMLResponse)
def invoice_home(request: Request):
    clients = db_list_clients(200)
    return jinja.TemplateResponse("invoice_home.html", {"request": request, "clients": clients})


@app.post("/invoice/client")
async def invoice_create_or_pick(request: Request):
    form = await request.form()
    cid_raw = (form.get("client_id") or "").strip()
    if cid_raw.isdigit():
        return RedirectResponse(url=f"/invoice/create?client_id={cid_raw}", status_code=303)

    fields = {
        "fio": (form.get("fio") or "").strip(),
        "passport": (form.get("passport") or "").strip(),
        "organ": (form.get("organ") or "").strip(),
        "vydan": normalize_date((form.get("vydan") or "").strip()),
        "address": (form.get("address") or "").strip(),
        "phone": (form.get("phone") or "").strip(),
        "contract_no": (form.get("contract_no") or "").strip(),
        "contract_date": normalize_date((form.get("contract_date") or "").strip()),
        "car_model": (form.get("car_model") or "").strip(),
        "customs_amount": (form.get("customs_amount") or "").strip(),
        "dkp_amount": (form.get("dkp_amount") or "").strip(),
    }
    cid = db_upsert_client(None, fields)
    return RedirectResponse(url=f"/invoice/create?client_id={cid}", status_code=303)


@app.get("/invoice/create", response_class=HTMLResponse)
def invoice_create_page(request: Request, client_id: int):
    c = db_get_client(client_id)
    if not c:
        return RedirectResponse(url="/invoice", status_code=303)
    return jinja.TemplateResponse(
        "invoice_create.html",
        {
            "request": request,
            "client": c,
            "purposes": INVOICE_PURPOSES,
        },
    )


@app.post("/invoice/generate")
async def invoice_generate(request: Request):
    try:
        if not INVOICE_TEMPLATE_PATH.exists():
            raise HTTPException(status_code=500, detail="Invoice template missing. Put template.docx into app/data or upload seed.")

        form = await request.form()
        cid_raw = (form.get("client_id") or "").strip()
        if not cid_raw.isdigit():
            raise HTTPException(status_code=400, detail="client_id is required")
        cid = int(cid_raw)
        client = db_get_client(cid)
        if not client:
            raise HTTPException(status_code=404, detail="Client not found")

        upd_fields = {
            "fio": (form.get("fio") or client.get("fio") or "").strip(),
            "passport": (form.get("passport") or client.get("passport") or "").strip(),
            "organ": (form.get("organ") or client.get("organ") or "").strip(),
            "vydan": normalize_date((form.get("vydan") or client.get("vydan") or "").strip()),
            "address": (form.get("address") or client.get("address") or "").strip(),
            "phone": (form.get("phone") or client.get("phone") or "").strip(),
            "contract_no": (form.get("contract_no") or client.get("contract_no") or "").strip(),
            "contract_date": normalize_date((form.get("contract_date") or client.get("contract_date") or "").strip()),
            "car_model": (form.get("car_model") or client.get("car_model") or "").strip(),
        }

        purpose_tpl = (form.get("purpose") or "").strip()
        if purpose_tpl not in INVOICE_PURPOSES:
            purpose_tpl = INVOICE_PURPOSES[0]
        purpose = purpose_tpl.format(модель=upd_fields.get("car_model") or "")

        amount_src = (form.get("amount") or "").strip()
        if not amount_src:
            raise HTTPException(status_code=400, detail="Amount is required")
        
        try:
            amount = format_amount_ru(amount_src)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Invalid amount: {str(e)}")

        if "тамож" in purpose_tpl.lower() or "пошлин" in purpose_tpl.lower():
            upd_fields["customs_amount"] = amount_src
        if "дкп" in purpose_tpl.lower():
            upd_fields["dkp_amount"] = amount_src

        db_upsert_client(cid, upd_fields)
        client = db_get_client(cid) or {**client, **upd_fields}

        req_mode = (form.get("req_mode") or "auto").strip()
        if req_mode == "auto":
            req_text = (form.get("req_text") or "").strip()
            if not req_text:
                raise HTTPException(status_code=400, detail="Requisites text is required")
            try:
                req = parse_requisites(req_text)
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Failed to parse requisites: {str(e)}")
        else:
            req = Requisites(
                recipient_fio=(form.get("recipient_fio") or "").strip(),
                bank_name=(form.get("bank_name") or "").strip(),
                bank_inn=normalize_digits((form.get("bank_inn") or "").strip()),
                bank_kpp=normalize_digits((form.get("bank_kpp") or "").strip()),
                bank_bik=normalize_digits((form.get("bank_bik") or "").strip()),
                bank_ks=group4((form.get("bank_ks") or "").strip()),
                bank_rs=group4((form.get("bank_rs") or "").strip()),
            )

        if not req.recipient_fio or not req.bank_rs:
            raise HTTPException(status_code=400, detail="Requisites empty/unrecognized. Please check recipient FIO and account number.")

        save_invoice_history(int(cid_raw), req, amount, purpose)

        mapping = {
            "CONTRACT_NO": (client.get("contract_no") or ""),
            "CONTRACT_DATE": (client.get("contract_date") or ""),
            "CLIENT_FIO": (client.get("fio") or ""),
            "CAR_MODEL": (client.get("car_model") or ""),
            "PAYMENT_PURPOSE": purpose,
            "AMOUNT": amount,
            "RECIPIENT_FIO": req.recipient_fio,
            "BANK_NAME": req.bank_name,
            "BANK_BIK": req.bank_bik,
            "BANK_INN": req.bank_inn,
            "BANK_KPP": req.bank_kpp,
            "BANK_KS": req.bank_ks,
            "BANK_RS": req.bank_rs,
        }

        out_id = uuid.uuid4().hex
        safe_name = safe_filename((client.get("fio") or "client"))
        out_docx = GENERATED_DIR / f"{out_id}_invoice_{safe_name}.docx"
        out_pdf = GENERATED_DIR / f"{out_id}_invoice_{safe_name}.pdf"

        # Проверяем существование шаблона
        if not INVOICE_TEMPLATE_PATH.exists():
            raise HTTPException(status_code=500, detail=f"Invoice template not found at {INVOICE_TEMPLATE_PATH}")

        try:
            # Копируем шаблон
            shutil.copy2(INVOICE_TEMPLATE_PATH, out_docx)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to copy template: {str(e)}")

        try:
            # Открываем и редактируем документ
            doc = Document(str(out_docx))
            apply_mapping(doc, mapping)
            doc.save(str(out_docx))
        except Exception as e:
            # Удаляем поврежденный файл
            if out_docx.exists():
                out_docx.unlink()
            raise HTTPException(status_code=500, detail=f"Failed to process document: {str(e)}")

        # Проверяем что файл действительно создан
        if not out_docx.exists():
            raise HTTPException(status_code=500, detail="Document was not created")

        # Сохраняем ссылку на документ в базе данных
        if cid:
            try:
                save_generated_document(cid, "invoice", out_id, f"{out_id}_invoice_{safe_name}.docx")
            except Exception:
                pass  # Не критично если не сохранилось в истории

        want_pdf = (form.get("want_pdf") or "").lower() in ("1", "true", "on", "yes")
        if want_pdf:
            try:
                with tempfile.TemporaryDirectory() as td:
                    tmp_docx = os.path.join(td, "invoice.docx")
                    tmp_pdf = os.path.join(td, "invoice.pdf")
                    shutil.copy2(out_docx, tmp_docx)
                    docx_to_pdf(tmp_docx, tmp_pdf)
                    shutil.copy2(tmp_pdf, out_pdf)
                return FileResponse(str(out_pdf), media_type="application/pdf", filename=f"Счет на оплату {safe_name}.pdf")
            except Exception as e:
                # Если PDF не удалось создать, возвращаем DOCX
                pass

        return FileResponse(
            str(out_docx),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"Счет на оплату {safe_name}.docx",
        )
    
    except HTTPException:
        raise
    except Exception as e:
        # Логируем ошибку и возвращаем более информативное сообщение
        import traceback
        error_details = traceback.format_exc()
        print(f"Error in invoice_generate: {error_details}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")


@app.get("/settings", response_class=HTMLResponse)
def settings_page(request: Request):
    s = load_settings()
    cur = get_currency_rates()
    return jinja.TemplateResponse("settings.html", {"request": request, "settings": s, "currency": cur})


@app.post("/settings")
async def settings_save_post(request: Request):
    form = await request.form()
    s = load_settings()
    cur = s.get("currency", {})
    cur["rub_to_byn"] = (form.get("rub_to_byn") or "").strip()
    cur["eur_byn"] = (form.get("eur_byn") or "").strip()
    cur["usd_byn"] = (form.get("usd_byn") or "").strip()
    s["currency"] = cur
    save_settings(s)
    return RedirectResponse(url="/settings", status_code=303)


@app.get("/tpo", response_class=HTMLResponse)
def tpo_home(request: Request):
    clients = db_list_clients(200)
    currency = get_currency_rates()
    return jinja.TemplateResponse("tpo_home.html", {"request": request, "clients": clients, "currency": currency})


@app.post("/tpo/client")
async def tpo_create_or_pick(request: Request):
    form = await request.form()
    cid_raw = (form.get("client_id") or "").strip()
    if cid_raw.isdigit():
        return RedirectResponse(url=f"/tpo/create?client_id={cid_raw}", status_code=303)

    fields = {
        "fio": (form.get("fio") or "").strip(),
        "passport": (form.get("passport") or "").strip(),
        "organ": (form.get("organ") or "").strip(),
        "vydan": normalize_date((form.get("vydan") or "").strip()),
        "address": (form.get("address") or "").strip(),
        "phone": (form.get("phone") or "").strip(),
        "car_model": (form.get("car_model") or "").strip(),
        "vin": (form.get("vin") or "").strip(),
        "obem": (form.get("obem") or "").strip(),
        "vypusk": normalize_date((form.get("vypusk") or "").strip()),
        "registr": normalize_date((form.get("registr") or "").strip()),
        "customs_amount": (form.get("cust_rub") or "").strip(),
        "dkp_amount": (form.get("price_rub") or "").strip(),
    }
    cid = db_upsert_client(None, fields)
    return RedirectResponse(url=f"/tpo/create?client_id={cid}", status_code=303)


@app.get("/tpo/create", response_class=HTMLResponse)
def tpo_create_page(request: Request, client_id: int):
    c = db_get_client(client_id)
    if not c:
        return RedirectResponse(url="/tpo", status_code=303)
    mapping: Dict[str, str] = {}
    try:
        mapping = json.loads(c.get("last_contract_json") or "{}")
    except Exception:
        mapping = {}

    prefill = {
        "fio": c.get("fio") or _pick_first(mapping, ["CLIENT_FIO", "FIO", "{{FIO}}", "ФИО"]),
        "address": c.get("address") or _pick_first(mapping, ["ADDRESS"]),
        "passport": c.get("passport") or _pick_first(mapping, ["PASSPORT"]),
        "vydan": c.get("vydan") or _pick_first(mapping, ["VYDAN", "DATA_PASS"]),
        "organ": c.get("organ") or _pick_first(mapping, ["ORGAN"]),
        "car_model": c.get("car_model") or _pick_first(mapping, ["CAR", "CAR_MODEL"]),
        "vin": c.get("vin") or _pick_first(mapping, ["VIN", "vin"]),
        "obem": c.get("obem") or _pick_first(mapping, ["OBEM"]),
        "vypusk": c.get("vypusk") or _pick_first(mapping, ["VYPUSK"]),
        "registr": c.get("registr") or _pick_first(mapping, ["REGISTR"]),
    }

    price_rub = (c.get("dkp_amount") or _pick_first(mapping, ["CAR_RUB"]) or "").strip()
    cust_rub = (c.get("customs_amount") or _pick_first(mapping, ["CUST_RUB"]) or "").strip()
    currency = get_currency_rates()
    return jinja.TemplateResponse(
        "tpo_create.html",
        {
            "request": request,
            "client": c,
            "prefill": prefill,
            "price_rub": price_rub,
            "cust_rub": cust_rub,
            "currency": currency,
        },
    )


@app.post("/tpo/generate")
async def tpo_generate(request: Request, qr: UploadFile = File(None)):
    if not TPO_TEMPLATE_PATH.exists():
        raise HTTPException(status_code=500, detail="TPO template missing. Put tpo_template.docx into app/data or seed /mnt/data/template.docx")

    form = await request.form()
    cid_raw = (form.get("client_id") or "").strip()
    if not cid_raw.isdigit():
        raise HTTPException(status_code=400, detail="client_id is required")
    cid = int(cid_raw)
    client = db_get_client(cid)
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")

    fio = (form.get("fio") or "").strip()
    address = (form.get("address") or "").strip()
    passport = (form.get("passport") or "").strip()
    vydan = normalize_date((form.get("vydan") or "").strip())
    organ = (form.get("organ") or "").strip()
    car = (form.get("car_model") or "").strip()
    vin = (form.get("vin") or "").strip()
    vypusk = normalize_date((form.get("vypusk") or "").strip())
    registr = normalize_date((form.get("registr") or "").strip())
    obem = (form.get("obem") or "").strip()
    price_rub = (form.get("price_rub") or "").strip()
    cust_rub = (form.get("cust_rub") or "").strip()

    rates = get_currency_rates()
    vals = calc_tpo_values(price_rub=price_rub, cust_rub=cust_rub, rates=rates)

    today = datetime.now().strftime("%d.%m.%Y")

    def cap(s: str) -> str:
        return (s or "").upper().strip()

    mapping = {
        "FIO": cap(fio),
        "ADDRESS": cap(address),
        "PASSPORT": cap(passport),
        "DATA_PASS": vydan,
        "VYDAN": vydan,
        "ORGAN": cap(organ),
        "CAR": cap(car),
        "VIN": cap(vin),
        "VYPUSK": vypusk,
        "REGISTR": registr,
        "OBEM": obem,
        "DATE": today,
        **vals,
    }

    upd_fields = {
        "fio": fio or (client.get("fio") or ""),
        "address": address or (client.get("address") or ""),
        "passport": passport or (client.get("passport") or ""),
        "vydan": vydan or (client.get("vydan") or ""),
        "organ": organ or (client.get("organ") or ""),
        "car_model": car or (client.get("car_model") or ""),
        "vin": vin or (client.get("vin") or ""),
        "obem": obem or (client.get("obem") or ""),
        "vypusk": vypusk or (client.get("vypusk") or ""),
        "registr": registr or (client.get("registr") or ""),
        "dkp_amount": price_rub,
        "customs_amount": cust_rub,
        "tpo_summa_byn": vals.get("CUST", ""),
        "tpo_price_eur": vals.get("PRICE_EUR", ""),
        "last_tpo_json": json.dumps(mapping, ensure_ascii=False),
    }
    db_upsert_client(cid, upd_fields)

    out_id = uuid.uuid4().hex
    safe_name = safe_filename(f"{fio}_{today}")
    out_docx = GENERATED_DIR / f"{out_id}_tpo_{safe_name}.docx"
    shutil.copy2(TPO_TEMPLATE_PATH, out_docx)
    doc = Document(str(out_docx))
    apply_mapping(doc, mapping)

    tmp_qr_path = None
    if qr and qr.filename:
        tmp_qr_path = str(GENERATED_DIR / f"{out_id}_qr_{safe_filename(qr.filename)}")
        with open(tmp_qr_path, "wb") as f:
            f.write(await qr.read())
        insert_qr_code(doc, tmp_qr_path, placeholder_key="QR CODE")

    doc.save(str(out_docx))

    # Сохраняем ссылку на документ в базе данных
    if cid:
        save_generated_document(cid, "tpo", out_id, f"{out_id}_tpo_{safe_name}.docx")

    if tmp_qr_path:
        try:
            os.remove(tmp_qr_path)
        except Exception:
            pass

    return FileResponse(
        str(out_docx),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"ТПО_{safe_name}.docx",
    )


@app.get("/t/{template_id}", response_class=HTMLResponse)
def fill_page(template_id: str, request: Request):
    t = get_template(template_id)
    clients = db_list_clients(200)
    
    placeholders = list(t.get("placeholders", []))
    
    for mandatory in MANDATORY_CONTRACT_FIELDS:
        if mandatory not in placeholders:
            placeholders.append(mandatory)
    
    fields = []
    for key in placeholders:
        preset = FIELD_PRESETS.get(key, {"label": key, "type": "text"})
        fields.append({"key": key, **preset})
    
    return jinja.TemplateResponse(
        "fill.html",
        {
            "request": request,
            "t": t,
            "fields": fields,
            "clients": clients,
        },
    )


@app.post("/generate/{template_id}")
async def generate(template_id: str, request: Request):
    t = get_template(template_id)
    form = await request.form()

    cid_raw = (form.get("client_id") or "").strip()
    cid = int(cid_raw) if cid_raw.isdigit() else None
    
    # Получаем данные клиента для подтягивания пустых полей
    client = db_get_client(cid) if cid else {}

    # ВАЖНО: Собираем все плейсхолдеры, включая обязательные
    placeholders = list(t.get("placeholders", []))
    for mandatory in MANDATORY_CONTRACT_FIELDS:
        if mandatory not in placeholders:
            placeholders.append(mandatory)

    mapping: Dict[str, str] = {}
    for key in placeholders:  # <-- Теперь обрабатываем ВСЕ поля, включая MANDATORY
        raw = (form.get(key) or "").strip()
        
        # Если поле пустое в форме, пытаемся взять из базы клиента
        if not raw and client:
            if key == "CONTRACT_DATE":
                raw = client.get("contract_date", "")
            elif key == "ORGAN":
                raw = client.get("organ", "")
            elif key == "VYDAN":
                raw = client.get("vydan", "")
            elif key == "CONTRACT_NO":
                raw = client.get("contract_no", "")
            elif key in ("CLIENT_FIO", "FIO"):
                raw = client.get("fio", "")
            elif key == "PASSPORT":
                raw = client.get("passport", "")
            elif key == "ADDRESS":
                raw = client.get("address", "")
            elif key == "PHONE":
                raw = client.get("phone", "")
        
        # Нормализуем даты
        if key in ("CONTRACT_DATE", "VYDAN") and raw:
            raw = normalize_date(raw)
        
        mapping[key] = raw

    client_fields = {
        "fio": _pick_first(mapping, ["CLIENT_FIO", "FIO", "ФИО"]),
        "passport": _pick_first(mapping, ["PASSPORT"]),
        "organ": _pick_first(mapping, ["ORGAN"]),
        "vydan": _pick_first(mapping, ["VYDAN"]),
        "address": _pick_first(mapping, ["ADDRESS"]),
        "phone": _pick_first(mapping, ["PHONE"]),
        "contract_no": _pick_first(mapping, ["CONTRACT_NO"]),
        "contract_date": _pick_first(mapping, ["CONTRACT_DATE"]),
        "car_model": _pick_first(mapping, ["CAR", "CAR_MODEL", "CAR_ALL"]),
        "vin": _pick_first(mapping, ["vin", "VIN"]),
        "obem": _pick_first(mapping, ["OBEM"]),
        "vypusk": _pick_first(mapping, ["year_car", "year car", "YEAR_CAR", "YEAR CAR", "VYPUSK", "ГОД", "YEAR"]),
        "registr": _pick_first(mapping, ["REGISTR"]),
        "color": _pick_first(mapping, ["color", "Color", "COLOR", "ЦВЕТ"]),
        "probeg": _pick_first(mapping, ["probeg", "Probeg", "PROBEG", "ПРОБЕГ"]),
        "engine": _pick_first(mapping, ["Engine", "ENGINE", "ДВИГАТЕЛЬ"]),
        "delivery": _pick_first(mapping, ["Delivery", "DELIVERY", "ДОСТАВКА"]),
        "customs_amount": _pick_first(mapping, ["CUST_RUB", "CUST"]) or _find_amount_like(mapping, ["CUST", "TAMOZH", "ПОШЛ", "DUTY"]),
        "dkp_amount": _pick_first(mapping, ["CAR_RUB", "DKP_AMOUNT"]) or _find_amount_like(mapping, ["DKP", "CAR_RUB", "PRICE", "СУММ", "СТОИ"]),
        "last_template_id": t.get("id", ""),
        "last_contract_json": json.dumps(mapping, ensure_ascii=False),
    }

    if any(v for v in client_fields.values() if v):
        cid = db_upsert_client(cid, client_fields)

    src = UPLOADS_DIR / t["filename"]
    if not src.exists():
        raise HTTPException(status_code=404, detail="Template file missing")

    out_id = uuid.uuid4().hex
    out_name_parts = [t.get("company", "template"), mapping.get("CONTRACT_NO", ""), mapping.get("CLIENT_FIO", "")]
    out_name = safe_filename("_".join([p for p in out_name_parts if p])) + ".docx"
    out_path = GENERATED_DIR / f"{out_id}_{out_name}"

    shutil.copy2(src, out_path)
    doc = Document(str(out_path))
    apply_mapping(doc, mapping)
    doc.save(str(out_path))
    
    # ВАЖНО: Сохраняем ссылку на созданный договор в профиле клиента
    if cid:
        try:
            save_generated_document(cid, "contract", out_id, f"{out_id}_{out_name}")
        except Exception:
            pass  # Не критично если не сохранилось

    return FileResponse(
        str(out_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=out_name,
    )


@app.get("/admin", response_class=HTMLResponse)
def admin(request: Request):
    templates = list_templates()
    return jinja.TemplateResponse("admin.html", {"request": request, "templates": templates})


@app.get("/admin/add", response_class=HTMLResponse)
def admin_add_get(request: Request):
    return jinja.TemplateResponse("add.html", {"request": request})


@app.post("/admin/add")
async def admin_add_post(
    company: str = Form(...),
    name: str = Form(...),
    file: UploadFile = File(...),
):
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Please upload a .docx file")

    template_id = uuid.uuid4().hex
    dest_name = f"{template_id}.docx"
    dest_path = UPLOADS_DIR / dest_name

    with dest_path.open("wb") as f:
        f.write(await file.read())

    try:
        doc = Document(str(dest_path))
        placeholders = extract_placeholders(doc)
    except Exception as e:
        dest_path.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail=f"Invalid .docx: {e}")

    db = load_db()
    db.setdefault("templates", []).append(
        {
            "id": template_id,
            "company": company.strip(),
            "name": name.strip(),
            "filename": dest_name,
            "placeholders": placeholders,
            "created_at": datetime.utcnow().isoformat() + "Z",
        }
    )
    save_db(db)

    return RedirectResponse(url=f"/t/{template_id}", status_code=303)


@app.post("/admin/delete/{template_id}")
def admin_delete(template_id: str):
    db = load_db()
    templates = db.get("templates", [])
    kept = [t for t in templates if t.get("id") != template_id]
    if len(kept) == len(templates):
        raise HTTPException(status_code=404, detail="Template not found")

    for t in templates:
        if t.get("id") == template_id:
            p = UPLOADS_DIR / t.get("filename", "")
            if p.exists():
                p.unlink()
            break

    db["templates"] = kept
    save_db(db)
    return RedirectResponse(url="/admin", status_code=303)


def summa_to_words_ru(summa_str: str) -> str:
    try:
        summa_clean = summa_str.replace(',', '.').replace(' ', '').replace('\u00A0', '')
        summa_float = float(summa_clean)
        rubles = int(summa_float)
        kopecks = int(round((summa_float - rubles) * 100))
        if kopecks == 100:
            rubles += 1
            kopecks = 0
        rubles_words = num2words(rubles, lang='ru')
        return f"{rubles_words} рублей {kopecks:02d} копеек"
    except:
        return summa_str


def replace_with_bold(doc: Document, mapping: Dict[str, str]) -> None:
    def process_paragraph(p):
        for key, value in mapping.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder not in p.text:
                continue
            
            full_text = p.text
            before, after = full_text.split(placeholder, 1)
            
            for r in p.runs:
                r.text = ""
            
            if before:
                p.add_run(before)
            
            run_value = p.add_run(value)
            run_value.bold = True
            
            if after:
                p.add_run(after)
            
            break
    
    for p in doc.paragraphs:
        process_paragraph(p)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)


@app.get("/kvit", response_class=HTMLResponse)
def kvit_home(request: Request):
    clients = db_list_clients(200)
    return jinja.TemplateResponse(
        "kvit_home.html",
        {
            "request": request,
            "clients": clients,
        },
    )


@app.post("/kvit/create")
async def kvit_create(request: Request):
    form = await request.form()
    
    cid_raw = (form.get("client_id") or "").strip()
    cid = int(cid_raw) if cid_raw.isdigit() else None
    client = db_get_client(cid) if cid else {}
    
    # Попытка получить данные из последнего договора
    contract_json = client.get("last_contract_json", "{}")
    try:
        contract_data = json.loads(contract_json)
    except:
        contract_data = {}
    
    # Подтягиваем данные с приоритетом: форма → договор → база данных
    fio = (form.get("fio") or "").strip() or contract_data.get("CLIENT_FIO", "") or contract_data.get("FIO", "") or client.get("fio", "")
    address = (form.get("address") or "").strip() or contract_data.get("ADDRESS", "") or client.get("address", "")
    passport = (form.get("passport") or "").strip() or contract_data.get("PASSPORT", "") or client.get("passport", "")
    vydan = (form.get("vydan") or "").strip() or contract_data.get("ORGAN", "") or client.get("organ", "")
    summa = (form.get("summa") or "").strip()
    
    if not summa and client.get("tpo_summa_byn"):
        summa = client.get("tpo_summa_byn", "")
    
    summa_formatted = summa.replace(".", ",")
    if "," not in summa_formatted:
        summa_formatted += ",00"
    
    mapping = {
        "CLIENT NAME": fio,
        "ADRESS": address,
        "PASSPORT": passport,
        "VYDAN": vydan,
        "SUMMA": summa_formatted,
    }
    
    upd_fields = {
        "fio": fio or (client.get("fio") or ""),
        "address": address or (client.get("address") or ""),
        "passport": passport or (client.get("passport") or ""),
        "organ": vydan or (client.get("organ") or ""),
    }
    db_upsert_client(cid, upd_fields)
    
    if not KVIT_TEMPLATE_PATH.exists():
        raise HTTPException(status_code=404, detail="Шаблон квитанции не найден")
    
    out_id = uuid.uuid4().hex
    today = datetime.now().strftime("%d.%m.%Y")
    safe_name = safe_filename(f"{fio}_{today}")
    out_docx = GENERATED_DIR / f"{out_id}_kvit_{safe_name}.docx"
    
    shutil.copy2(KVIT_TEMPLATE_PATH, out_docx)
    doc = Document(str(out_docx))
    replace_with_bold(doc, mapping)
    doc.save(str(out_docx))
    
    return FileResponse(
        str(out_docx),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"Квитанция_{safe_name}.docx",
    )


@app.get("/util", response_class=HTMLResponse)
def util_home(request: Request):
    clients = db_list_clients(200)
    companies_data = load_companies()
    companies_list = companies_data.get("companies", [])
    
    return jinja.TemplateResponse(
        "util_home.html",
        {
            "request": request,
            "clients": clients,
            "companies": companies_list,
            "companies_json": json.dumps(companies_list, ensure_ascii=False),
        },
    )


@app.post("/util/create")
async def util_create(request: Request):
    form = await request.form()
    
    cid_raw = (form.get("client_id") or "").strip()
    cid = int(cid_raw) if cid_raw.isdigit() else None
    client = db_get_client(cid) if cid else {}
    
    # Попытка получить данные из последнего договора
    contract_json = client.get("last_contract_json", "{}")
    try:
        contract_data = json.loads(contract_json)
    except:
        contract_data = {}
    
    # Получаем company_id из договора (last_template_id)
    last_template_id = client.get("last_template_id", "")
    template_company = ""
    template_company_inn = ""
    template_company_address = ""
    
    if last_template_id:
        # Получаем данные компании из шаблона договора
        try:
            template = get_template(last_template_id)
            company_name = template.get("company", "")
            
            # Ищем компанию в companies.json
            companies_data = load_companies()
            for comp in companies_data.get("companies", []):
                if comp.get("name") == company_name or comp.get("id") == company_name.lower():
                    template_company = f'ООО "{comp.get("name", "")}"'
                    template_company_inn = comp.get("inn", "")
                    template_company_address = comp.get("address", "")
                    break
        except:
            pass
    
    # Подтягиваем данные с приоритетом: форма → договор → шаблон → база данных
    fio = (form.get("fio") or "").strip() or contract_data.get("CLIENT_FIO", "") or contract_data.get("FIO", "") or client.get("fio", "")
    passport = (form.get("passport") or "").strip() or contract_data.get("PASSPORT", "") or client.get("passport", "")
    address = (form.get("address") or "").strip() or contract_data.get("ADDRESS", "") or client.get("address", "")
    company = (form.get("company") or "").strip() or template_company or client.get("company_name", "")
    inn = (form.get("inn") or "").strip() or template_company_inn or client.get("company_inn", "")
    company_address = (form.get("company_address") or "").strip() or template_company_address or client.get("company_address", "")
    
    # Данные авто - подтягиваем из договора и базы
    car = (form.get("car") or "").strip() or contract_data.get("CAR", "") or contract_data.get("CAR_MODEL", "") or client.get("car_model", "")
    year_car_full = (form.get("year_car") or "").strip() or contract_data.get("year_car", "") or contract_data.get("year car", "") or contract_data.get("YEAR_CAR", "") or contract_data.get("VYPUSK", "") or client.get("vypusk", "")
    
    # Извлекаем только год (ГГГГ) из возможных форматов: "2020", "01.2020", "15.01.2020"
    year_car = year_car_full
    if year_car_full:
        # Если формат дд.мм.гггг или мм.гггг - извлекаем только год
        parts = year_car_full.split('.')
        if len(parts) > 1:
            year_car = parts[-1]  # Последняя часть - это год
        # Проверяем, что это 4 цифры
        if year_car and year_car.isdigit() and len(year_car) == 4:
            year_car = year_car
        else:
            year_car = year_car_full
    
    color = contract_data.get("color", "") or contract_data.get("Color", "") or contract_data.get("COLOR", "") or client.get("color", "")
    probeg = contract_data.get("probeg", "") or contract_data.get("Probeg", "") or contract_data.get("PROBEG", "") or client.get("probeg", "")
    vin = (form.get("vin") or "").strip() or contract_data.get("VIN", "") or client.get("vin", "")
    
    # Склонения ФИО
    fio1 = fio  # Именительный падеж - как есть: Иванов Иван Иванович
    fio_genitive = decline_fio_genitive(fio)  # Родительный: Иванова Ивана Ивановича
    fio_instrumental = decline_fio_instrumental(fio)  # Творительный: Ивановым Иваном Ивановичем
    
    nomer = (form.get("nomer") or "").strip()
    bank = (form.get("bank") or "").strip()
    account = (form.get("account") or "").strip()
    summa_rub = (form.get("summa_rub") or "").strip()
    
    now = datetime.now()
    date_number = now.strftime("%d")
    months_ru = {
        1: "января", 2: "февраля", 3: "марта", 4: "апреля",
        5: "мая", 6: "июня", 7: "июля", 8: "августа",
        9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
    }
    date_month = months_ru.get(now.month, now.strftime("%B"))
    year = now.strftime("%Y")
    
    summa_text = summa_to_words_ru(summa_rub)
    
    mapping = {
        "fio": fio_genitive,  # Родительный: Иванова Ивана Ивановича
        "1fio": fio_instrumental,  # Творительный: Ивановым Иваном Ивановичем
        "fio1": fio1,  # Именительный: Иванов Иван Иванович
        "passport": passport,
        "address": company_address,  # Юридический адрес компании, а не клиента
        "company": company,
        "inn": inn,
        "car": car,
        "year car": year_car,  # Только год ГГГГ
        "year_car": year_car,  # Только год ГГГГ
        "color": color,
        "probeg": probeg,
        "vin": vin,
        "nomer": nomer,
        "bank": bank,
        "account": account,
        "summa": summa_rub,
        "summa text": summa_text,
        "date number": date_number,
        "date mouth": date_month,
        "year": year,
    }
    
    # Получаем company_id из формы (если выбрана компания из списка)
    company_id = (form.get("company_id") or "").strip()
    
    upd_fields = {
        "fio": fio or (client.get("fio") or ""),
        "passport": passport or (client.get("passport") or ""),
        "address": address or (client.get("address") or ""),
        "car_model": car or (client.get("car_model") or ""),
        "vin": vin or (client.get("vin") or ""),
        "vypusk": year_car or (client.get("vypusk") or ""),
        "color": color or (client.get("color") or ""),
        "probeg": probeg or (client.get("probeg") or ""),
        "company_inn": inn or (client.get("company_inn") or ""),
        "company_address": company_address or (client.get("company_address") or ""),
        "company_name": company or (client.get("company_name") or ""),
        "company_id": company_id or (client.get("company_id") or ""),
    }
    db_upsert_client(cid, upd_fields)
    
    if not UTIL_TEMPLATE_PATH.exists():
        raise HTTPException(status_code=404, detail="Шаблон письма не найден")
    
    out_id = uuid.uuid4().hex
    today = datetime.now().strftime("%d.%m.%Y")
    safe_name = safe_filename(f"{fio}_{today}")
    out_docx = GENERATED_DIR / f"{out_id}_util_{safe_name}.docx"
    
    shutil.copy2(UTIL_TEMPLATE_PATH, out_docx)
    doc = Document(str(out_docx))
    apply_mapping(doc, mapping)
    doc.save(str(out_docx))
    
    return FileResponse(
        str(out_docx),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"Письмо_утильсбор_{safe_name}.docx",
    )


@app.get("/dkp", response_class=HTMLResponse)
def dkp_home(request: Request):
    clients = db_list_clients(200)
    companies_data = load_companies()
    company_names = [c.get("name", "") for c in companies_data.get("companies", [])]
    return jinja.TemplateResponse(
        "dkp_home.html",
        {
            "request": request,
            "clients": clients,
            "companies": company_names,
        },
    )


@app.post("/dkp/client")
async def dkp_select_client(request: Request):
    form = await request.form()
    cid_raw = (form.get("client_id") or "").strip()
    
    if cid_raw.isdigit():
        return RedirectResponse(url=f"/dkp/create?client_id={cid_raw}", status_code=303)
    
    fields = {
        "fio": (form.get("fio") or "").strip(),
        "passport": (form.get("passport") or "").strip(),
        "address": (form.get("address") or "").strip(),
        "phone": (form.get("phone") or "").strip(),
        "car_model": "",
        "vin": "",
    }
    new_id = db_upsert_client(None, fields)
    return RedirectResponse(url=f"/dkp/create?client_id={new_id}", status_code=303)


@app.get("/dkp/create", response_class=HTMLResponse)
def dkp_create_form(request: Request, client_id: int = 0):
    client = db_get_client(client_id) if client_id else {}
    if not client:
        return RedirectResponse(url="/dkp", status_code=303)

    # Источники данных для подтягивания (в порядке приоритета):
    # 1. Данные клиента из профиля
    # 2. История счетов (для ФИО получателя платежа)
    # 3. Последний договор (last_contract_json)
    # 4. Данные ТПО (last_tpo_json)

    # Пытаемся получить данные из последнего договора
    contract_json = client.get("last_contract_json", "{}")
    try:
        contract_data = json.loads(contract_json)
    except:
        contract_data = {}

    # Пытаемся получить данные из ТПО
    tpo_json = client.get("last_tpo_json", "{}")
    try:
        tpo_data = json.loads(tpo_json)
    except:
        tpo_data = {}

    # Получаем историю счетов для поиска ФИО получателя платежа
    invoice_hist = get_invoice_history(client_id, limit=1)
    invoice_data = invoice_hist[0] if invoice_hist else {}

    # ФИО - пробуем из разных источников:
    # 1. ФИО клиента из профиля
    # 2. ФИО получателя платежа из счета
    # 3. ФИО из последнего договора
    fio = (
        client.get("fio", "") or
        invoice_data.get("recipient_fio", "") or
        contract_data.get("CLIENT_FIO", "") or
        contract_data.get("FIO", "") or
        tpo_data.get("FIO", "") or
        ""
    )

    # Паспорт - пробуем из разных источников
    passport = (
        client.get("passport", "") or
        contract_data.get("PASSPORT", "") or
        tpo_data.get("PASSPORT", "") or
        ""
    )

    # Адрес - пробуем из разных источников
    address = (
        client.get("address", "") or
        contract_data.get("ADDRESS", "") or
        tpo_data.get("ADDRESS", "") or
        ""
    )

    # Телефон - пробуем из разных источников
    phone = (
        client.get("phone", "") or
        contract_data.get("PHONE", "") or
        ""
    )

    # Автомобиль - пробуем из разных источников
    car = (
        client.get("car_model", "") or
        contract_data.get("CAR", "") or
        contract_data.get("CAR_MODEL", "") or
        tpo_data.get("CAR", "") or
        ""
    )

    # VIN - пробуем из разных источников
    vin = (
        client.get("vin", "") or
        contract_data.get("VIN", "") or
        tpo_data.get("VIN", "") or
        ""
    )

    # Объём двигателя
    obem = (
        client.get("obem", "") or
        contract_data.get("OBEM", "") or
        tpo_data.get("OBEM", "") or
        ""
    )

    # EUR цена из ТПО - ключевой момент!
    tpo_price_eur = client.get("tpo_price_eur", "")

    # Если цена в EUR не найдена в TPO, пробуем извлечь из DKP суммы
    if not tpo_price_eur:
        dkp_amount = client.get("dkp_amount", "")
        if dkp_amount:
            try:
                # Простая конвертация из рублей в EUR по курсу
                rates = get_currency_rates()
                eur_byn = float(rates.get("eur_byn", 4.0))
                rub_to_byn = float(rates.get("rub_to_byn", 0.04))
                price_rub = float(dkp_amount.replace(" ", "").replace(",", "."))
                price_byn = price_rub * rub_to_byn
                tpo_price_eur = f"{price_byn / eur_byn:.2f}"
            except (ValueError, TypeError):
                tpo_price_eur = ""

    # Если всё ещё нет цены в EUR, ищем в данных договора
    if not tpo_price_eur:
        tpo_price_eur = (
            contract_data.get("PRICE_EUR", "") or
            contract_data.get("price_eur", "") or
            contract_data.get("PRICE_EURO", "") or
            contract_data.get("CAR_PRICE_EUR", "") or
            tpo_data.get("PRICE_EUR", "") or
            ""
        )

    # Цвет - переводим на немецкий, если на русском
    color_raw = (
        client.get("color", "") or
        contract_data.get("Color", "") or
        contract_data.get("COLOR", "") or
        contract_data.get("color", "") or
        tpo_data.get("Color", "") or
        ""
    )
    # Переводим цвет на немецкий
    color = translate_color_ru_to_de(color_raw)
    
    # Если перевод пустой, используем исходное значение
    if not color and color_raw:
        color = color_raw

    # Год выпуска - приоритет: данные клиента -> данные договора -> данные ТПО
    year = (
        client.get("vypusk", "") or
        contract_data.get("year_car", "") or  # с нижним подчеркиванием
        contract_data.get("year car", "") or  # с пробелом
        contract_data.get("YEAR_CAR", "") or
        contract_data.get("YEAR CAR", "") or
        contract_data.get("YEAR", "") or
        contract_data.get("ГОД", "") or
        tpo_data.get("VYPUSK", "") or
        ""
    )

    # Пробег - приоритет: данные клиента -> данные договора -> данные ТПО
    probeg = (
        client.get("probeg", "") or
        contract_data.get("Probeg", "") or
        contract_data.get("PROBEG", "") or
        contract_data.get("probeg", "") or
        contract_data.get("ПРОБЕГ", "") or
        tpo_data.get("Probeg", "") or
        ""
    )

    # Цена в EUR из ТПО
    price_eur = (
        tpo_price_eur or
        contract_data.get("PRICE_EUR", "") or
        contract_data.get("price_eur", "") or
        tpo_data.get("PRICE_EUR", "") or
        ""
    )

    # Первая регистрация (мм.гггг) - преобразуем из дд.мм.гггг
    regdate_raw = (
        client.get("registr", "") or
        contract_data.get("REGISTR", "") or
        tpo_data.get("REGISTR", "") or
        ""
    )
    # Преобразуем дату в формат мм.гггг
    regdate = convert_date_to_mmyyyy(regdate_raw)

    # Мощность в л.с.
    ps = (
        contract_data.get("PS", "") or
        tpo_data.get("PS", "") or
        ""
    )

    prefill = {
        "fio": fio,
        "fio_translit": translit_ru_to_lat(fio),
        "passport": format_passport_dkp(passport),
        "address": address,
        "phone": format_phone_dkp(phone),
        "car": car,
        "vin": format_vin_dkp(vin),
        "obem": obem,
        "year": year,
        "color": color,
        "probeg": probeg,
        "price_eur": price_eur,
        "regdate": regdate,
        "ps": ps,
    }

    companies_data = load_companies()
    company_names = [c.get("name", "") for c in companies_data.get("companies", [])]
    
    return jinja.TemplateResponse(
        "dkp_create.html",
        {
            "request": request,
            "client": client,
            "prefill": prefill,
            "companies": company_names,
        },
    )


@app.post("/dkp/generate")
async def dkp_generate(request: Request):
    form = await request.form()
    
    cid_raw = (form.get("client_id") or "").strip()
    cid = int(cid_raw) if cid_raw.isdigit() else None
    client = db_get_client(cid) if cid else {}
    
    # Получаем данные из ТПО для правильного заполнения Price и регистрации
    tpo_json = client.get("last_tpo_json", "{}")
    try:
        tpo_data = json.loads(tpo_json)
    except:
        tpo_data = {}
    
    date = (form.get("date") or datetime.now().strftime("%d.%m.%Y")).strip()
    fio_ru = (form.get("fio") or "").strip()
    fio = translit_ru_to_lat(fio_ru)
    passport = format_passport_dkp((form.get("passport") or "").strip())
    address_ru = (form.get("address") or "").strip()
    address = translate_ru_to_de(address_ru)
    phone = format_phone_dkp((form.get("phone") or "").strip())
    company_ru = (form.get("company") or "").strip()
    company = translit_ru_to_lat(company_ru)  # Транслитерация названия компании
    car = (form.get("car") or "").strip()
    vin = format_vin_dkp((form.get("vin") or "").strip())
    obem = (form.get("obem") or "").strip()
    engine = (form.get("engine") or "Benzin").strip()
    
    # Регистрация - преобразуем из дд.мм.гггг в мм.гггг
    regdate_raw = (form.get("regdate") or "").strip()
    # Если пусто, берем из ТПО
    if not regdate_raw:
        regdate_raw = tpo_data.get("REGISTR", "")
    regdate = convert_date_to_mmyyyy(regdate_raw) if regdate_raw else format_regdate_dkp(regdate_raw)
    
    # Цена - берется из PRICE_EUR (из ТПО)
    price_eur_raw = (form.get("price_eur") or "").strip()
    # Если пусто, берем из ТПО
    if not price_eur_raw:
        price_eur_raw = (
            tpo_data.get("PRICE_EUR", "") or 
            tpo_data.get("price_eur", "") or 
            client.get("tpo_price_eur", "")
        )
    price_eur = price_eur_raw
    price = price_eur  # Price = PRICE_EUR
    
    # prtotal = PRICE_EUR + 50, округленное до ближайшего 10
    try:
        price_value = float(price_eur.replace(",", ".").replace(" ", ""))
        prtotal_raw = price_value + 50
        # Округляем до ближайшего 10
        prtotal_rounded = round(prtotal_raw / 10) * 10
        prtotal = f"{prtotal_rounded:.2f}".replace(".", ",")
    except (ValueError, TypeError, AttributeError):
        prtotal = price_eur
    
    # Получаем данные из договора для подстановки
    contract_json = client.get("last_contract_json", "{}")
    try:
        contract_data = json.loads(contract_json)
    except:
        contract_data = {}
    
    # Год выпуска - подтягиваем из разных источников
    year = (form.get("year") or "").strip()
    if not year:
        year = (
            client.get("vypusk", "") or
            contract_data.get("year_car", "") or
            contract_data.get("year car", "") or
            contract_data.get("YEAR_CAR", "") or
            contract_data.get("YEAR CAR", "") or
            contract_data.get("YEAR", "") or
            tpo_data.get("VYPUSK", "") or
            ""
        )
    
    # Цвет - автоматически переводим на немецкий
    color_input = (form.get("color") or "").strip()
    if not color_input:
        color_input = (
            client.get("color", "") or
            contract_data.get("Color", "") or
            contract_data.get("COLOR", "") or
            contract_data.get("color", "") or
            tpo_data.get("Color", "") or
            ""
        )
    
    # Переводим цвет на немецкий
    color_de = translate_color_ru_to_de(color_input)
    
    # Если перевод пустой, используем исходное значение
    if not color_de and color_input:
        color_de = color_input
    
    color_ru = color_input if color_input.lower() in COLOR_TRANSLATION else ""
    
    probeg = (form.get("probeg") or "").strip()
    ps = (form.get("ps") or "").strip()
    kw = ps_to_kw(ps) if ps else ""
    gear = (form.get("gear") or "Automatik").strip()
    seat = (form.get("seat") or "5").strip()
    
    mapping = {
        "date": date,
        "fio": fio,
        "passport": passport,
        "address": address,
        "phone": phone,
        "company": company,  # Транслитерированная компания
        "car": car,
        "vin": vin,
        "obem": obem,
        "engine": engine,
        "REGISTR": regdate,  # мм.гггг формат - REGISTR вместо VYPUSK
        "VYPUSK": regdate,  # Для совместимости со старыми шаблонами
        "regdate": regdate,  # оставляем для совместимости
        "Price": price,  # PRICE_EUR из ТПО
        "price": price,
        "price_eur": price_eur,
        "PRICE_EUR": price_eur,  # С большими буквами и подчеркиванием
        "prtotal": prtotal,  # PRICE_EUR + 50
        # Год выпуска - все варианты
        "year": year,
        "year_car": year,
        "year car": year,
        "YEAR": year,
        "YEAR_CAR": year,
        "YEAR CAR": year,
        # Цвет на немецком - все варианты
        "Color": color_de,
        "color": color_de,
        "COLOR": color_de,
        "Probeg": probeg,
        "probeg": probeg,
        "PROBEG": probeg,
        "ps": ps,
        "kw": kw,
        "gear": gear,
        "seat": seat,
    }
    
    # Создаем словарь с правильными ключами для подтягивания в форму ДКП
    contract_save_data = {
        "CLIENT_FIO": fio_ru,
        "PASSPORT": passport,
        "ADDRESS": address_ru,
        "PHONE": phone,
        "CAR": car,
        "CAR_MODEL": car,
        "VIN": vin,
        "OBEM": obem,
        "YEAR": year,
        "YEAR CAR": year,
        "year car": year,
        "YEAR_CAR": year,
        "year_car": year,
        "Color": color_de,  # сохраняем цвет на НЕМЕЦКОМ
        "COLOR": color_de,
        "color": color_de,
        "Probeg": probeg,
        "PROBEG": probeg,
        "probeg": probeg,
        "PRICE_EUR": price_eur,
        "price_eur": price_eur,
        "Price": price,
        "PS": ps,
        "REGISTR": regdate,
        "VYPUSK": regdate,
    }
    
    upd_fields = {
        "fio": fio_ru or (client.get("fio") or ""),
        "passport": passport or (client.get("passport") or ""),
        "address": address_ru or (client.get("address") or ""),
        "phone": phone or (client.get("phone") or ""),
        "car_model": car or (client.get("car_model") or ""),
        "vin": vin or (client.get("vin") or ""),
        "obem": obem or (client.get("obem") or ""),
        "vypusk": year or (client.get("vypusk") or ""),
        "color": color_de or (client.get("color") or ""),  # сохраняем цвет на НЕМЕЦКОМ
        "probeg": probeg or (client.get("probeg") or ""),
        "registr": regdate or (client.get("registr") or ""),
        "tpo_price_eur": price_eur or (client.get("tpo_price_eur") or ""),
        "last_contract_json": json.dumps(contract_save_data, ensure_ascii=False),
        "dkp_amount": price or client.get("dkp_amount", ""),
    }
    if cid:
        db_upsert_client(cid, upd_fields)
    
    if not DKP_TEMPLATE_PATH.exists():
        raise HTTPException(status_code=404, detail="Шаблон ДКП не найден. Убедитесь что файл dkp_template.docx находится в папке data/")
    
    out_id = uuid.uuid4().hex
    today = datetime.now().strftime("%d.%m.%Y")
    safe_name = safe_filename(f"DKP_{fio}_{today}")
    out_docx = GENERATED_DIR / f"{out_id}_{safe_name}.docx"
    
    shutil.copy2(DKP_TEMPLATE_PATH, out_docx)
    doc = Document(str(out_docx))
    apply_mapping(doc, mapping)
    doc.save(str(out_docx))
    
    # Сохраняем ссылку на документ в базе данных
    if cid:
        save_generated_document(cid, "dkp", out_id, f"{out_id}_{safe_name}.docx")
    
    return FileResponse(
        str(out_docx),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"Autoverkaufsvertrag_{safe_name}.docx",
    )


@app.get("/api/translate")
def api_translate(text: str = "", lang: str = "de"):
    if not text:
        return {"translated": ""}
    
    if lang == "de":
        translated = translate_ru_to_de(text)
    else:
        translated = text
    
    return {"translated": translated}


@app.get("/api/translit")
def api_translit(text: str = ""):
    if not text:
        return {"translit": ""}
    
    return {"translit": translit_ru_to_lat(text)}


def load_companies() -> Dict[str, Any]:
    if not COMPANIES_PATH.exists():
        return {"companies": [], "settings": {}}
    try:
        with open(COMPANIES_PATH, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return {"companies": [], "settings": {}}


def save_companies(data: Dict[str, Any]) -> None:
    with open(COMPANIES_PATH, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def get_company(company_id: str) -> Optional[Dict[str, Any]]:
    data = load_companies()
    for company in data.get("companies", []):
        if company.get("id") == company_id:
            return company
    return None


def list_companies() -> List[Dict[str, Any]]:
    return load_companies().get("companies", [])


def db_migrate_extended() -> None:
    con = db_connect()
    cur = con.cursor()

    # Создаем таблицу invoice_history
    cur.execute("""
        CREATE TABLE IF NOT EXISTS invoice_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            client_id INTEGER NOT NULL,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            recipient_fio TEXT DEFAULT '',
            bank_name TEXT DEFAULT '',
            bank_inn TEXT DEFAULT '',
            bank_kpp TEXT DEFAULT '',
            bank_bik TEXT DEFAULT '',
            bank_ks TEXT DEFAULT '',
            bank_rs TEXT DEFAULT '',
            amount TEXT DEFAULT '',
            purpose TEXT DEFAULT '',
            FOREIGN KEY (client_id) REFERENCES clients(id)
        )
    """)

    # Создаем таблицу proxy
    cur.execute("""
        CREATE TABLE IF NOT EXISTS proxy (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            client_id INTEGER NOT NULL,
            company_id TEXT DEFAULT '',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            proxy_date TEXT DEFAULT '',
            srok TEXT DEFAULT '',
            fio TEXT DEFAULT '',
            passport TEXT DEFAULT '',
            address TEXT DEFAULT '',
            car_model TEXT DEFAULT '',
            vin TEXT DEFAULT '',
            generated_filename TEXT DEFAULT '',
            FOREIGN KEY (client_id) REFERENCES clients(id)
        )
    """)

    # Создаем таблицу generated_documents для хранения ссылок на документы
    # Сначала проверяем, существует ли таблица
    cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='generated_documents'")
    if cur.fetchone() is None:
        # Таблица не существует, создаем её
        cur.execute("""
            CREATE TABLE generated_documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                client_id INTEGER NOT NULL,
                doc_type TEXT NOT NULL DEFAULT '',
                doc_id TEXT DEFAULT '',
                filename TEXT DEFAULT '',
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (client_id) REFERENCES clients(id)
            )
        """)
        print("Создана таблица generated_documents")
    else:
        # Проверяем структуру таблицы
        cur.execute("PRAGMA table_info(generated_documents)")
        cols = [r[1] for r in cur.fetchall()]
        # Если нет колонки doc_type, пересоздаем таблицу
        if 'doc_type' not in cols:
            cur.execute("DROP TABLE generated_documents")
            cur.execute("""
                CREATE TABLE generated_documents (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    client_id INTEGER NOT NULL,
                    doc_type TEXT NOT NULL DEFAULT '',
                    doc_id TEXT DEFAULT '',
                    filename TEXT DEFAULT '',
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (client_id) REFERENCES clients(id)
                )
            """)
            print("Пересоздана таблица generated_documents")

    # Проверяем необходимость добавления новых колонок в таблицу клиентов
    cur.execute("PRAGMA table_info(clients)")
    existing = {r[1] for r in cur.fetchall()}

    new_columns = [
        ("last_proxy_fio", "TEXT DEFAULT ''"),
        ("last_company_id", "TEXT DEFAULT ''"),
        ("tpo_price_eur", "TEXT DEFAULT ''"),
        ("color", "TEXT DEFAULT ''"),
        ("probeg", "TEXT DEFAULT ''"),
    ]

    for col, decl in new_columns:
        if col not in existing:
            cur.execute(f"ALTER TABLE clients ADD COLUMN {col} {decl}")

    con.commit()
    con.close()


db_migrate_extended()


def save_invoice_history(client_id: int, req: Requisites, amount: str, purpose: str) -> int:
    con = db_connect()
    cur = con.cursor()
    cur.execute("""
        INSERT INTO invoice_history 
        (client_id, recipient_fio, bank_name, bank_inn, bank_kpp, bank_bik, bank_ks, bank_rs, amount, purpose)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        client_id, req.recipient_fio, req.bank_name, req.bank_inn, req.bank_kpp,
        req.bank_bik, req.bank_ks, req.bank_rs, amount, purpose
    ))
    con.commit()
    new_id = int(cur.lastrowid)
    con.close()
    return new_id


def get_invoice_history(client_id: int, limit: int = 50) -> List[Dict[str, Any]]:
    con = db_connect()
    cur = con.cursor()
    cur.execute("""
        SELECT * FROM invoice_history 
        WHERE client_id = ? 
        ORDER BY id DESC LIMIT ?
    """, (client_id, limit))
    rows = cur.fetchall()
    con.close()
    return [dict(r) for r in rows]


def save_generated_document(client_id: int, doc_type: str, doc_id: str, filename: str) -> int:
    """Сохранить ссылку на сгенерированный документ"""
    con = db_connect()
    cur = con.cursor()
    cur.execute("""
        INSERT INTO generated_documents (client_id, doc_type, doc_id, filename)
        VALUES (?, ?, ?, ?)
    """, (client_id, doc_type, doc_id, filename))
    con.commit()
    new_id = int(cur.lastrowid)
    con.close()
    return new_id


def get_client_documents(client_id: int, doc_type: Optional[str] = None) -> List[Dict[str, Any]]:
    """Получить список сгенерированных документов клиента"""
    con = db_connect()
    cur = con.cursor()
    
    if doc_type:
        cur.execute("""
            SELECT * FROM generated_documents 
            WHERE client_id = ? AND doc_type = ?
            ORDER BY id DESC
        """, (client_id, doc_type))
    else:
        cur.execute("""
            SELECT * FROM generated_documents 
            WHERE client_id = ?
            ORDER BY id DESC
        """, (client_id,))
    
    rows = cur.fetchall()
    con.close()
    return [dict(r) for r in rows]


def get_document_by_id(doc_id: str) -> Optional[Dict[str, Any]]:
    """Получить информацию о документе по его ID"""
    con = db_connect()
    cur = con.cursor()
    cur.execute("SELECT * FROM generated_documents WHERE doc_id = ?", (doc_id,))
    row = cur.fetchone()
    con.close()
    return dict(row) if row else None


@app.get("/proxy", response_class=HTMLResponse)
def proxy_home(request: Request):
    clients = db_list_clients(200)
    companies = list_companies()
    companies_data = load_companies()
    default_srok = companies_data.get("settings", {}).get("default_proxy_srok", "25 мая 2026 г.")
    default_date = companies_data.get("settings", {}).get("default_proxy_date", "25 мая 2025 г.")
    return jinja.TemplateResponse(
        "proxy_home.html",
        {
            "request": request,
            "clients": clients,
            "companies": companies,
            "default_srok": default_srok,
            "default_date": default_date,
        },
    )


@app.post("/proxy/client")
async def proxy_create_or_pick(request: Request):
    form = await request.form()
    cid_raw = (form.get("client_id") or "").strip()
    
    if cid_raw.isdigit():
        return RedirectResponse(url=f"/proxy/create?client_id={cid_raw}", status_code=303)
    
    fields = {
        "fio": (form.get("fio") or "").strip(),
        "passport": (form.get("passport") or "").strip(),
        "address": (form.get("address") or "").strip(),
    }
    new_id = db_upsert_client(None, fields)
    return RedirectResponse(url=f"/proxy/create?client_id={new_id}", status_code=303)


@app.get("/proxy/create", response_class=HTMLResponse)
def proxy_create_form(request: Request, client_id: int):
    client = db_get_client(client_id)
    if not client:
        return RedirectResponse(url="/proxy", status_code=303)
    
    companies = list_companies()
    companies_data = load_companies()
    default_srok = companies_data.get("settings", {}).get("default_proxy_srok", "25 мая 2026 г.")
    default_date = companies_data.get("settings", {}).get("default_proxy_date", "25 мая 2025 г.")
    
    # ВАЖНО: Доверенность делается на ПОЛУЧАТЕЛЯ ПЛАТЕЖА из реквизитов счета, а не на клиента!
    # Получаем ФИО получателя платежа из истории счетов
    invoice_hist = get_invoice_history(client_id, limit=10)  # Последние 10 получателей
    requisites_fio = ""
    recent_recipients = []
    
    if invoice_hist and len(invoice_hist) > 0:
        # Первый получатель - для автоподстановки
        invoice = invoice_hist[0]
        requisites_fio = invoice.get("recipient_fio", "") or ""
        
        # Собираем уникальных получателей из истории
        seen = set()
        for inv in invoice_hist:
            fio = inv.get("recipient_fio", "")
            if fio and fio not in seen:
                recent_recipients.append(fio)
                seen.add(fio)
    
    # Для доверенности используем:
    # - ФИО: из реквизитов получателя платежа (если есть), иначе пустое поле
    # - Паспорт, орган, дата: НЕ берем из данных клиента, это данные другого человека (получателя платежа)
    # - Автомобиль: из данных клиента (это информация о машине, а не о человеке)
    prefill = {
        "fio": requisites_fio,  # ФИО поверенного (получателя платежа)
        "passport": "",  # Паспорт поверенного - заполняется вручную
        "organ": "",  # Орган выдачи - заполняется вручную
        "vydan": "",  # Дата выдачи - заполняется вручную
        "address": "",  # Адрес поверенного - опционально
        "car_model": client.get("car_model", ""),
        "vin": client.get("vin", ""),
    }
    
    # Передаем информацию о том, что ФИО взято из реквизитов
    has_requisites_fio = bool(requisites_fio)
    
    return jinja.TemplateResponse(
        "proxy_create.html",
        {
            "request": request,
            "client": client,
            "prefill": prefill,
            "companies": companies,
            "default_srok": default_srok,
            "default_date": default_date,
            "has_requisites_fio": has_requisites_fio,
            "requisites_fio": requisites_fio,
            "recent_recipients": recent_recipients,  # Список последних получателей
        },
    )


@app.post("/proxy/generate")
async def proxy_generate(request: Request):
    if not PROXY_TEMPLATE_PATH.exists():
        raise HTTPException(status_code=500, detail="Proxy template missing. Put proxy_template.docx into app/data/")
    
    form = await request.form()
    
    cid_raw = (form.get("client_id") or "").strip()
    cid = int(cid_raw) if cid_raw.isdigit() else None
    client = db_get_client(cid) if cid else {}
    
    company_id = (form.get("company_id") or "").strip()
    company = get_company(company_id) if company_id else {}
    
    proxy_date = (form.get("proxy_date") or "").strip()
    srok = (form.get("srok") or "").strip()
    fio = (form.get("fio") or "").strip().upper()
    passport = (form.get("passport") or "").strip()
    organ = (form.get("organ") or "").strip()  # Кем выдан паспорт
    vydan_raw = (form.get("vydan") or "").strip()  # Дата выдачи паспорта
    address = (form.get("address") or "").strip()
    car_model = (form.get("car_model") or "").strip()
    vin = (form.get("vin") or "").strip()
    
    # Форматируем дату выдачи в формат дд.мм.гггг
    vydan = vydan_raw
    if vydan_raw and '-' in vydan_raw:
        # Преобразуем из формата гггг-мм-дд в дд.мм.гггг
        parts = vydan_raw.split('-')
        if len(parts) == 3:
            vydan = f"{parts[2]}.{parts[1]}.{parts[0]}"
    
    # Парсим реквизиты если предоставлены
    req_mode = (form.get("req_mode") or "").strip()
    if req_mode == "auto":
        req_text = (form.get("req_text") or "").strip()
        if req_text:
            req = parse_requisites(req_text)
            if req.recipient_fio:
                fio = req.recipient_fio.upper()
    
    # Форматируем паспорт для шаблона (формат XX XX XXXXXX)
    passport_formatted = passport
    if passport:
        import re
        digits = re.sub(r'\D', '', passport)
        if len(digits) >= 10:
            passport_formatted = f"{digits[:2]} {digits[2:4]} {digits[4:10]}"
        elif len(digits) >= 4:
            passport_formatted = f"{digits[:2]} {digits[2:]}"

    # Парсим дату выдачи паспорта из адреса или используем текущую
    passport_issue = ""
    passport_date = ""
    import re
    # Ищем дату в формате ДД.ММ.ГГГГ или ГГГГ-ММ-ДД
    date_match = re.search(r'(\d{1,2})[\.\-](\d{1,2})[\.\-](\d{2,4})', address)
    if date_match:
        day, month, year = date_match.groups()
        if len(year) == 2:
            year = "20" + year if int(year) < 50 else "19" + year
        passport_date = f"{day.zfill(2)}.{month.zfill(2)}.{year}"
        # Удаляем дату из адреса
        address_clean = re.sub(r'\d{1,2}[\.\-]\d{1,2}[\.\-]\d{2,4}', '', address).strip()
        address = address_clean

    # Формируем номер документа (если передан из формы, используем его, иначе генерируем рандомный)
    doc_number = (form.get("doc_number") or "").strip()
    if not doc_number:
        # Генерируем номер формата XXYYYY (например: 65А78, 12В34)
        import random
        letters = "АВЕКМНОРСТУХ"
        doc_number = f"{random.randint(10, 99)}{random.choice(letters)}{random.randint(1000, 9999)}"

    # Маппинг для шаблона - используем правильные имена плейсхолдеров из шаблона
    mapping = {
        "NOMER": f"№{doc_number}",
        "Date": proxy_date,
        "City": company.get("city", ""),
        "SROK": srok,
        "RECIPIENT_FIO": fio,
        "RECIPIENT": fio,
        "FIO": fio,
        "fio": fio,
        # Паспорт - разные варианты плейсхолдеров
        "PASSPORT": passport_formatted,
        "Passport": passport_formatted,
        "passport": passport_formatted,
        "Drop_Pass": passport_formatted,
        "DROP_PASSPORT": passport_formatted,
        # Адрес - разные варианты
        "ADDRESS": address,
        "Address": address,
        "address": address,
        "Drop_Address": address,
        "DROP_ADDRESS": address,
        # Орган выдачи паспорта и дата выдачи
        "DROP_ORGAN": organ,  # Кем выдан паспорт
        "ORGAN": organ,
        "DROP_VYDAN": vydan,  # Дата выдачи паспорта
        "VYDAN": vydan,
        "DATA_PASS": vydan,
        # Автомобиль
        "CAR": car_model,
        "Car": car_model,
        "car": car_model,
        # VIN
        "VIN": vin,
        "Vin": vin,
        "vin": vin,
        # Компания
        "COMPANY": company.get("name", ""),
        "Company": company.get("name", ""),
        "company": company.get("name", ""),
        # Город компании
        "COMPANY_CITY": company.get("city", ""),
        "CompanyCity": company.get("city", ""),
        "City": company.get("city", ""),
        "city": company.get("city", ""),
        # Адрес компании
        "COMPANY_ADDRESS": company.get("address", ""),
        "CompanyAddress": company.get("address", ""),
        "Comp_Address": company.get("address", ""),
        # Директор
        "DIRECTOR_FIO": company.get("director_fio", ""),
        "DirectorFIO": company.get("director_fio", ""),
        "director_fio": company.get("director_fio", ""),
        "Director": company.get("director_fio", ""),
        "director": company.get("director_fio", ""),
        "DIRECTOR": company.get("director_fio", ""),
    }
    
    # ВАЖНО: Сохраняем только данные АВТОМОБИЛЯ и последнего поверенного
    # НЕ сохраняем ФИО и паспорт поверенного в профиль клиента!
    # Доверенность на ПОЛУЧАТЕЛЯ ПЛАТЕЖА, а не на клиента
    upd_fields = {
        # Данные авто - обновляем
        "car_model": car_model or (client.get("car_model") or ""),
        "vin": vin or (client.get("vin") or ""),
        # История - кто был последним поверенным
        "last_proxy_fio": fio,
        "last_company_id": company_id,
    }
    if cid:
        db_upsert_client(cid, upd_fields)
    
    # Сохраняем в историю прокси
    if cid:
        con = db_connect()
        cur = con.cursor()
        cur.execute("""
            INSERT INTO proxy 
            (client_id, company_id, proxy_date, srok, fio, passport, address, car_model, vin)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (cid, company_id, proxy_date, srok, fio, passport, address, car_model, vin))
        con.commit()
        proxy_id = int(cur.lastrowid)
        con.close()
    
    out_id = uuid.uuid4().hex
    today = datetime.now().strftime("%d.%m.%Y")
    safe_name = safe_filename(f"{fio}_{today}")
    out_docx = GENERATED_DIR / f"{out_id}_proxy_{safe_name}.docx"
    
    shutil.copy2(PROXY_TEMPLATE_PATH, out_docx)
    doc = Document(str(out_docx))
    apply_mapping(doc, mapping)
    doc.save(str(out_docx))
    
    # Сохраняем ссылку на документ в базе данных
    if cid:
        save_generated_document(cid, "proxy", out_id, f"{out_id}_proxy_{safe_name}.docx")
    
    return FileResponse(
        str(out_docx),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"Доверенность_{safe_name}.docx",
    )


@app.get("/clients", response_class=HTMLResponse)
def clients_list(request: Request):
    """Страница со списком всех клиентов"""
    search = request.query_params.get("search", "").strip().lower()
    clients = db_list_clients(500)
    
    if search:
        clients = [c for c in clients if search in c.get("fio", "").lower() or search in c.get("phone", "").lower() or search in c.get("contract_no", "").lower()]
    
    return jinja.TemplateResponse(
        "clients.html",
        {
            "request": request,
            "clients": clients,
            "search": search,
        },
    )


@app.get("/client/{client_id}", response_class=HTMLResponse)
def client_profile(request: Request, client_id: int):
    client = db_get_client(client_id)
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    invoice_hist = get_invoice_history(client_id)
    
    # Получаем все сгенерированные документы клиента
    all_documents = get_client_documents(client_id)
    
    # Группируем документы по типам
    invoices = [d for d in all_documents if d.get("doc_type") == "invoice"]
    tpos = [d for d in all_documents if d.get("doc_type") == "tpo"]
    dkps = [d for d in all_documents if d.get("doc_type") == "dkp"]
    proxies = [d for d in all_documents if d.get("doc_type") == "proxy"]
    
    contracts = []
    try:
        con = db_connect()
        cur = con.cursor()
        cur.execute("""
            SELECT * FROM contracts WHERE client_id = ? ORDER BY id DESC LIMIT 20
        """, (client_id,))
        for r in cur.fetchall():
            contracts.append(dict(r))
        con.close()
    except Exception:
        pass
    
    proxy_hist = []
    try:
        con = db_connect()
        cur = con.cursor()
        cur.execute("""
            SELECT * FROM proxy WHERE client_id = ? ORDER BY id DESC LIMIT 20
        """, (client_id,))
        for r in cur.fetchall():
            proxy_hist.append(dict(r))
        con.close()
    except Exception:
        pass
    
    tpo_price_eur = client.get("tpo_price_eur", "")
    
    return jinja.TemplateResponse(
        "client_profile.html",
        {
            "request": request,
            "client": client,
            "invoice_history": invoice_hist,
            "contracts": contracts,
            "proxy_history": proxy_hist,
            "tpo_price_eur": tpo_price_eur,
            "invoices": invoices,
            "tpos": tpos,
            "dkps": dkps,
            "generated_proxies": proxies,
        },
    )


@app.post("/client/{client_id}/update")
async def client_update(request: Request, client_id: int):
    client = db_get_client(client_id)
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    form = await request.form()
    
    upd_fields = {
        "fio": (form.get("fio") or "").strip(),
        "passport": (form.get("passport") or "").strip(),
        "organ": (form.get("organ") or "").strip(),
        "vydan": normalize_date((form.get("vydan") or "").strip()),
        "address": (form.get("address") or "").strip(),
        "phone": (form.get("phone") or "").strip(),
        "contract_no": (form.get("contract_no") or "").strip(),
        "contract_date": normalize_date((form.get("contract_date") or "").strip()),
        "car_model": (form.get("car_model") or "").strip(),
        "vin": (form.get("vin") or "").strip(),
        "obem": (form.get("obem") or "").strip(),
        "vypusk": (form.get("vypusk") or "").strip(),
        "color": (form.get("color") or "").strip(),
        "probeg": (form.get("probeg") or "").strip(),
        "customs_amount": (form.get("customs_amount") or "").strip(),
        "dkp_amount": (form.get("dkp_amount") or "").strip(),
        "company_inn": (form.get("company_inn") or "").strip(),
        "company_address": (form.get("company_address") or "").strip(),
    }
    
    db_upsert_client(client_id, upd_fields)
    return RedirectResponse(url=f"/client/{client_id}", status_code=303)


@app.post("/client/{client_id}/delete")
def client_delete(client_id: int):
    """Удалить клиента и все его данные"""
    client = db_get_client(client_id)
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    db_delete_client(client_id)
    return RedirectResponse(url="/clients", status_code=303)


@app.post("/client/{client_id}/clear-history")
def client_clear_history(client_id: int):
    """Очистить историю документов клиента"""
    client = db_get_client(client_id)
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    db_clear_client_history(client_id)
    return RedirectResponse(url=f"/client/{client_id}", status_code=303)


@app.get("/api/client/{client_id}/requisites")
def api_client_requisites(client_id: int):
    history = get_invoice_history(client_id, limit=1)
    if history:
        return {"requisites": history[0]}
    return {"requisites": None}


@app.get("/api/client/{client_id}/last-proxy")
def api_client_last_proxy(client_id: int):
    client = db_get_client(client_id)
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    return {
        "fio": client.get("last_proxy_fio", ""),
        "company_id": client.get("last_company_id", ""),
    }


@app.get("/api/client/{client_id}/documents")
def api_client_documents(client_id: int, doc_type: Optional[str] = None):
    """Получить список документов клиента"""
    client = db_get_client(client_id)
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    documents = get_client_documents(client_id, doc_type)
    return {"documents": documents}


@app.get("/api/client/{client_id}/invoice-history")
def api_client_invoices(client_id: int):
    """Получить историю счетов клиента"""
    client = db_get_client(client_id)
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    invoices = get_invoice_history(client_id)
    return {"invoices": invoices}


@app.get("/document/{doc_id}")
def view_document(doc_id: str):
    """Просмотр сгенерированного документа"""
    doc_info = get_document_by_id(doc_id)
    if not doc_info:
        raise HTTPException(status_code=404, detail="Document not found")
    
    filename = doc_info.get("filename", "")
    filepath = GENERATED_DIR / filename
    
    if not filepath.exists():
        raise HTTPException(status_code=404, detail="Document file not found")
    
    # Определяем тип документа для правильного content-type
    if filename.endswith('.pdf'):
        media_type = "application/pdf"
    else:
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    return FileResponse(
        str(filepath),
        media_type=media_type,
        filename=filename,
    )


@app.get("/admin/companies", response_class=HTMLResponse)
def companies_admin(request: Request):
    companies = list_companies()
    companies_data = load_companies()
    settings = companies_data.get("settings", {})
    return jinja.TemplateResponse(
        "companies_admin.html",
        {
            "request": request,
            "companies": companies,
            "settings": settings,
        },
    )


@app.post("/admin/companies/add")
async def companies_add(request: Request):
    form = await request.form()
    
    name = (form.get("name") or "").strip()
    city = (form.get("city") or "").strip()
    director_fio = (form.get("director_fio") or "").strip()
    address = (form.get("address") or "").strip()
    
    if not name:
        raise HTTPException(status_code=400, detail="Company name is required")
    
    companies_data = load_companies()
    
    company_id = re.sub(r'[^a-zа-я0-9]', '', name.lower().replace(' ', '_'))
    
    new_company = {
        "id": company_id,
        "name": name,
        "city": city,
        "director_fio": director_fio,
        "address": address,
        "inn": "",
        "ogrn": ""
    }
    
    companies_data.setdefault("companies", []).append(new_company)
    save_companies(companies_data)
    
    return RedirectResponse(url="/admin/companies", status_code=303)


@app.post("/admin/companies/delete/{company_id}")
def companies_delete(company_id: str):
    companies_data = load_companies()
    companies_data["companies"] = [c for c in companies_data.get("companies", []) if c.get("id") != company_id]
    save_companies(companies_data)
    return RedirectResponse(url="/admin/companies", status_code=303)


@app.post("/admin/companies/settings")
async def companies_settings(request: Request):
    form = await request.form()
    
    companies_data = load_companies()
    companies_data.setdefault("settings", {})["default_proxy_date"] = (form.get("default_proxy_date") or "").strip()
    companies_data["settings"]["default_proxy_srok"] = (form.get("default_proxy_srok") or "").strip()
    
    save_companies(companies_data)
    return RedirectResponse(url="/admin/companies", status_code=303)
